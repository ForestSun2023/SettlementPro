import sys
import time
import io
import copy
import os
import json
import urllib.request
import urllib.error
import tempfile
import numpy as np
from datetime import datetime, timedelta
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                             QHBoxLayout, QTableWidget, QTableWidgetItem, QPushButton,
                             QLabel, QLineEdit, QHeaderView, QMessageBox, QGroupBox, QTextEdit,
                             QFileDialog, QCheckBox, QSplashScreen, QComboBox, QColorDialog, QDialog, QGridLayout, QDoubleSpinBox, QDateEdit, QInputDialog, QScrollArea, QToolButton, QFrame)
from PyQt6.QtGui import QPixmap, QPainter, QColor, QFont, QIcon, QPen
from PyQt6.QtCore import Qt, QDate
import matplotlib.dates as mdates

# 尝试导入可选依赖库
try:
    from openpyxl import Workbook
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import matplotlib.pyplot as plt
from matplotlib.patches import Patch
from mpl_toolkits.axes_grid1.inset_locator import inset_axes

# 设置中文字体，防止绘图乱码
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
plt.rcParams['axes.unicode_minus'] = False

# --- 暗黑模式样式表 ---
DARK_STYLESHEET = """
QMainWindow, QWidget {
    background-color: #2b2b2b;
    color: #ffffff;
}
QGroupBox {
    border: 1px solid #555;
    border-radius: 5px;
    margin-top: 10px;
    font-weight: bold;
}
QGroupBox::title {
    subcontrol-origin: margin;
    subcontrol-position: top center;
    padding: 0 5px;
    color: #ffffff;
}
QTableWidget {
    background-color: #3c3f41;
    color: #ffffff;
    gridline-color: #555;
    selection-background-color: #4b6eaf;
}
QHeaderView::section {
    background-color: #3c3f41;
    color: #ffffff;
    border: 1px solid #555;
}
QLineEdit, QTextEdit {
    background-color: #3c3f41;
    color: #ffffff;
    border: 1px solid #555;
    border-radius: 3px;
}
QPushButton {
    background-color: #3c3f41;
    color: #ffffff;
    border: 1px solid #555;
    padding: 6px;
    border-radius: 4px;
}
QPushButton:hover {
    background-color: #4c5052;
}
QPushButton:pressed {
    background-color: #2c2f31;
}
QCheckBox {
    color: #ffffff;
}
QLabel {
    color: #ffffff;
}
"""

class GreyModelGM11:
    """
    GM(1,1) 灰色预测模型核心算法类
    """
    def __init__(self):
        self.a = None  # 发展系数
        self.b = None  # 灰作用量
        self.x0 = None # 原始序列
        self.x1 = None # 累加序列
        self.fitted = False

    def fit(self, data):
        """训练模型"""
        try:
            self.x0 = np.array(data, dtype=float)
            n = len(self.x0)
            if n < 4:
                return False, "数据量太少，至少需要4期数据"

            # 1. 一次累加生成 (1-AGO)
            self.x1 = np.cumsum(self.x0)

            # 2. 构造数据矩阵 B 和数据向量 Y
            B = []
            Y = []
            for i in range(1, n):
                # 紧邻均值生成
                z = 0.5 * (self.x1[i] + self.x1[i-1])
                B.append([-z, 1])
                Y.append(self.x0[i])
            
            B = np.array(B)
            Y = np.array(Y).reshape(-1, 1)

            # 3. 最小二乘法求解
            # (B^T * B)^-1 * B^T * Y
            bt_b = np.dot(B.T, B)
            if np.linalg.det(bt_b) == 0:
                return False, "矩阵奇异，无法求解"
                
            bt_b_inv = np.linalg.inv(bt_b)
            params = np.dot(np.dot(bt_b_inv, B.T), Y)
            
            self.a = params[0][0]
            self.b = params[1][0]
            self.fitted = True
            return True, "模型构建成功"
        except Exception as e:
            return False, str(e)

    def predict(self, k_times):
        """预测未来 k_times 期 (包含已有的期数)"""
        if not self.fitted:
            return None

        predict_x1 = [] 
        predict_x0 = [] 

        # 响应方程
        x0_1 = self.x0[0]
        term1 = x0_1 - self.b / self.a
        
        for k in range(k_times):
            if k == 0:
                val = x0_1
                predict_x1.append(val)
                predict_x0.append(val)
            else:
                # 计算累加预测值
                val = term1 * np.exp(-self.a * k) + self.b / self.a
                predict_x1.append(val)
                # 累减还原得到原始预测值
                predict_x0.append(val - predict_x1[-2])
        
        return np.array(predict_x0)

    def get_accuracy(self, predicted_vals):
        """计算残差和相对误差"""
        n = len(self.x0)
        residuals = self.x0 - predicted_vals[:n]
        relative_error = np.abs(residuals) / self.x0 * 100
        return residuals, relative_error

    def residual_check(self, relative_errors, threshold=10.0):
        """
        残差检验 (精度等级评定)
        :param relative_errors: 相对误差百分比数组
        :param threshold: 合格阈值 (默认 10%)
        :return: (bool 是否合格, str 评价信息)
        """
        avg_err = np.mean(relative_errors)
        if avg_err < 1.0: grade = "一级 (优)"
        elif avg_err < 5.0: grade = "二级 (合格)"
        elif avg_err < 10.0: grade = "三级 (勉强)"
        else: grade = "四级 (不合格)"
            
        msg = f"精度评级: {grade}\n平均相对误差: {avg_err:.2f}%"
        return avg_err <= threshold, msg

    def posterior_variance_check(self, residuals):
        """
        后验差检验 (C值/P值) - 更严格的精度评定
        :param residuals: 残差数组
        :return: (bool 是否合格, str 评价信息, float C, float P)
        """
        n = len(residuals)
        if n < 1: return False, "数据不足", 0.0, 0.0
        
        # S1: 原始数据标准差
        s1 = np.std(self.x0)
        
        # S2: 残差标准差
        s2 = np.std(residuals)
        
        # C值: 后验差比值 (越小越好)
        C = s2 / s1 if s1 != 0 else 0.0
        
        # P值: 小误差概率 p = P{|e(k)-e_bar| < 0.6745*S1} (越大越好)
        res_mean = np.mean(residuals)
        threshold = 0.6745 * s1
        p_count = np.sum(np.abs(residuals - res_mean) < threshold)
        P = p_count / n
        
        # 评级标准 (参照工程测量规范)
        if P > 0.95 and C < 0.35: grade = "一级 (好)"; qualified = True
        elif P > 0.80 and C < 0.50: grade = "二级 (合格)"; qualified = True
        elif P > 0.70 and C < 0.65: grade = "三级 (勉强)"; qualified = True
        else: grade = "四级 (不合格)"; qualified = False
            
        msg = f"精度评级: {grade}\n后验差比值 C = {C:.4f}\n小误差概率 P = {P:.4f}"
        return qualified, msg, C, P

    def rolling_predict_check(self, data, min_len=4):
        """
        滚动预测检验：模拟每增加一期数据就重新建模预测下一期
        :param data: 实测数据序列
        :param min_len: 起始建模的最少数据量 (默认4)
        :return: 滚动预测值序列 (与 data 等长, 前 min_len 个为原始值)
        """
        n = len(data)
        # 初始化结果数组，默认填充原始值（因为前几期无法滚动预测）
        rolling_results = np.array(data, dtype=float)
        
        # 从第 min_len 期开始 (索引 min_len)，预测该期
        # 例如 min_len=4，我们有 0,1,2,3 四个数，第一次预测索引 4 的数
        for i in range(min_len, n):
            # 使用 [0, i-1] 的数据 (共 i 个点) 来预测第 i 个点
            history = data[:i]
            
            temp_model = GreyModelGM11()
            success, _ = temp_model.fit(history)
            
            if success:
                # 预测 i+1 期 (即包含当前点 i)
                # predict 返回长度为 len(history)+1 的数组，最后一个即为预测值
                preds = temp_model.predict(len(history) + 1)
                rolling_results[i] = preds[-1]
            else:
                rolling_results[i] = np.nan # 建模失败标记
                
        return rolling_results

class SimpleKalmanFilter:
    """
    简易一维卡尔曼滤波器
    用于平滑沉降观测数据，去除随机噪声，提取真实趋势
    """
    def __init__(self, Q=0.01, R=1.0):
        self.Q = Q  # 过程噪声协方差 (Process Noise) - 假设系统状态变化的波动
        self.R = R  # 观测噪声协方差 (Measurement Noise) - 仪器的测量误差
        self.x_hat = None # 后验估计值 (最佳估计)
        self.P = 1.0      # 后验误差协方差

    def filter(self, data):
        results = []
        for z in data:
            if self.x_hat is None:
                self.x_hat = z
                results.append(z)
                continue
            
            # 1. 预测 (Time Update): 假设静态或缓变过程 x(k) = x(k-1)
            x_pred = self.x_hat
            P_pred = self.P + self.Q
            
            # 2. 更新 (Measurement Update)
            K = P_pred / (P_pred + self.R) # 卡尔曼增益
            self.x_hat = x_pred + K * (z - x_pred)
            self.P = (1 - K) * P_pred
            
            results.append(self.x_hat)
        return np.array(results)

class MovingAverageFilter:
    """
    移动平均滤波器 (Moving Average Filter)
    通过计算滑动窗口内的平均值来平滑数据，消除短期波动
    """
    def __init__(self, window_size=3):
        self.window_size = window_size

    def filter(self, data):
        data = np.array(data)
        result = []
        for i in range(len(data)):
            # 窗口范围：从 i-window_size+1 到 i (包含 i)
            start_idx = max(0, i - self.window_size + 1)
            window = data[start_idx : i + 1]
            result.append(np.mean(window))
        return np.array(result)

class ExponentialMovingAverageFilter:
    """
    指数移动平均滤波器 (EMA)
    给予近期数据更高权重，对趋势变化反应更灵敏
    """
    def __init__(self, alpha=0.3):
        self.alpha = alpha  # 平滑系数 (0 < alpha < 1)，值越大越接近实测值

    def filter(self, data):
        data = np.array(data)
        if len(data) == 0: return np.array([])
        result = [data[0]] # 第一项直接作为初始值
        for i in range(1, len(data)):
            # EMA_today = alpha * value_today + (1-alpha) * EMA_yesterday
            val = self.alpha * data[i] + (1 - self.alpha) * result[-1]
            result.append(val)
        return np.array(result)

class StyleConfigDialog(QDialog):
    """图表样式配置对话框"""
    def __init__(self, current_styles, parent=None):
        super().__init__(parent)
        self.setWindowTitle("图表样式配置")
        self.resize(400, 300)
        # 深拷贝当前样式，避免直接修改
        self.temp_styles = copy.deepcopy(current_styles)
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        grid = QGridLayout()
        
        headers = ["图层名称", "颜色", "线型", "线宽"]
        for col, h in enumerate(headers):
            grid.addWidget(QLabel(h), 0, col)

        # 映射内部键名到显示名称
        self.names_map = {
            'observed': '实测值',
            'predicted': '预测值',
            'poly': '多项式拟合',
            'rolling': '滚动预测',
            'kalman': '卡尔曼滤波',
            'ma': '移动平均',
            'ema': '指数平滑',
            'conf_int': '置信区间 (填充)'
        }
        
        row = 1
        for key, name in self.names_map.items():
            style = self.temp_styles[key]
            
            grid.addWidget(QLabel(name), row, 0)
            
            # 颜色按钮
            btn_color = QPushButton()
            btn_color.setFixedSize(50, 25)
            self.set_btn_color(btn_color, style['color'])
            # 使用 lambda 捕获变量
            btn_color.clicked.connect(lambda checked, k=key, b=btn_color: self.choose_color(k, b))
            grid.addWidget(btn_color, row, 1)
            
            # 线型下拉框
            combo_style = QComboBox()
            combo_style.addItems(['-', '--', '-.', ':'])
            combo_style.setCurrentText(style['style'])
            combo_style.currentTextChanged.connect(lambda t, k=key: self.update_style(k, 'style', t))
            grid.addWidget(combo_style, row, 2)
            
            # 线宽微调框
            spin_width = QDoubleSpinBox()
            spin_width.setRange(0.5, 10.0)
            spin_width.setSingleStep(0.5)
            spin_width.setValue(style['width'])
            spin_width.valueChanged.connect(lambda v, k=key: self.update_style(k, 'width', v))
            grid.addWidget(spin_width, row, 3)
            
            row += 1

        layout.addLayout(grid)
        
        btn_box = QHBoxLayout()
        btn_ok = QPushButton("确定")
        btn_ok.clicked.connect(self.accept)
        btn_cancel = QPushButton("取消")
        btn_cancel.clicked.connect(self.reject)
        btn_box.addWidget(btn_ok)
        btn_box.addWidget(btn_cancel)
        layout.addLayout(btn_box)
        
        self.setLayout(layout)

    def set_btn_color(self, btn, color_str):
        btn.setStyleSheet(f"background-color: {color_str}; border: 1px solid #555;")
        
    def choose_color(self, key, btn):
        c = QColorDialog.getColor(QColor(self.temp_styles[key]['color']), self, "选择颜色")
        if c.isValid():
            hex_c = c.name()
            self.temp_styles[key]['color'] = hex_c
            self.set_btn_color(btn, hex_c)

    def update_style(self, key, field, value):
        self.temp_styles[key][field] = value
        
    def get_styles(self):
        return self.temp_styles

class ApiSettingsDialog(QDialog):
    """API 设置对话框"""
    def __init__(self, settings, parent=None):
        super().__init__(parent)
        self.setWindowTitle("API 设置")
        self.resize(420, 220)
        self.settings = settings
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout()
        form = QGridLayout()

        form.addWidget(QLabel("API 地址:"), 0, 0)
        self.input_endpoint = QLineEdit(self.settings.get("endpoint", ""))
        form.addWidget(self.input_endpoint, 0, 1)

        form.addWidget(QLabel("模型名:"), 1, 0)
        self.input_model = QLineEdit(self.settings.get("model", ""))
        form.addWidget(self.input_model, 1, 1)

        form.addWidget(QLabel("API Key:"), 2, 0)
        self.input_key = QLineEdit(self.settings.get("api_key", ""))
        self.input_key.setEchoMode(QLineEdit.EchoMode.Password)
        form.addWidget(self.input_key, 2, 1)

        layout.addLayout(form)

        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("保存")
        btn_cancel = QPushButton("取消")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        self.setLayout(layout)

    def get_settings(self):
        return {
            "endpoint": self.input_endpoint.text().strip(),
            "model": self.input_model.text().strip(),
            "api_key": self.input_key.text().strip()
        }

class MplCanvas(FigureCanvas):
    """Matplotlib 绘图画布"""
    def __init__(self, parent=None, width=5, height=4, dpi=100):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        # --- 修改：创建两个子图 (上图显示沉降，下图显示速率) ---
        gs = self.fig.add_gridspec(2, 1, height_ratios=[3, 1], hspace=0.08)
        self.ax1 = self.fig.add_subplot(gs[0]) # 主图
        self.ax2 = self.fig.add_subplot(gs[1], sharex=self.ax1) # 子图 (共享X轴)
        super(MplCanvas, self).__init__(self.fig)

class CollapsibleBox(QWidget):
    """简易可折叠面板"""
    def __init__(self, title="", parent=None):
        super().__init__(parent)
        self.toggle_button = QToolButton(text=title, checkable=True, checked=True)
        self.toggle_button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextBesideIcon)
        self.toggle_button.setArrowType(Qt.ArrowType.DownArrow)
        self.toggle_button.setStyleSheet("QToolButton { border: none; font-weight: bold; }")
        self.toggle_button.toggled.connect(self.on_toggled)

        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(4, 4, 4, 0)
        header_layout.addWidget(self.toggle_button)
        header_layout.addStretch()

        self.content_area = QFrame()
        self.content_area.setFrameShape(QFrame.Shape.StyledPanel)
        self.content_area.setFrameShadow(QFrame.Shadow.Plain)
        self.content_layout = QVBoxLayout()
        self.content_layout.setContentsMargins(8, 6, 8, 8)
        self.content_area.setLayout(self.content_layout)

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(2, 2, 2, 2)
        main_layout.addLayout(header_layout)
        main_layout.addWidget(self.content_area)

    def setContentLayout(self, layout):
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            widget = item.widget()
            if widget:
                widget.setParent(None)
        self.content_layout.addLayout(layout)

    def on_toggled(self, checked):
        self.content_area.setVisible(checked)
        self.toggle_button.setArrowType(Qt.ArrowType.DownArrow if checked else Qt.ArrowType.RightArrow)

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("沉降监测与变形预测系统 - 测绘工程课设")
        self.resize(1100, 700)
        
        self.gm_model = GreyModelGM11()
        self.is_dark_mode = False # 记录当前主题状态
        self.points_data = {}
        self.points_stage = {}
        self.point_positions = {}
        self.current_point = None
        self.is_loading_point = False
        self.last_raw_data = None
        self.last_predicted = None
        self.last_residuals = None
        self.last_rel_errors = None
        self.last_model_name = None
        self.last_model_params = None
        self.last_model_rmse = None
        self.last_model_mae = None
        self.multi_ref_points = []
        self.last_ref_stability_text = None
        self.last_change_index = None
        self.last_stage_stats = None
        self.last_stage_segments = None
        self.last_outlier_indices = []
        self.last_risk_text = None
        self.last_stability_text = None
        self.last_unstable_refs = []
        self.last_corrected = False
        self.point_results = {}
        self.last_deform_grades = None
        self.last_stage_deep = None
        self.multi_point_analysis = None
        self.api_endpoint = "https://api.deepseek.com/chat/completions"
        self.api_model = "deepseek-chat"
        self.api_key = ""
        self.load_api_settings()
        
        # --- 初始化图表样式配置 ---
        self.plot_styles = {
            'observed': {'color': '#0000FF', 'style': '-', 'width': 1.5, 'marker': 'o'},
            'predicted': {'color': '#FF0000', 'style': '--', 'width': 1.5, 'marker': '*'},
            'poly': {'color': '#008000', 'style': '--', 'width': 1.5, 'marker': '^'},
            'rolling': {'color': '#FF00FF', 'style': '-.', 'width': 1.5, 'marker': '.'},
            'kalman': {'color': '#00FFFF', 'style': '-', 'width': 2.0, 'marker': None},
            'ma': {'color': '#FFA500', 'style': '-', 'width': 1.5, 'marker': '.'},
            'ema': {'color': '#9C27B0', 'style': '-', 'width': 2.0, 'marker': None},
            'conf_int': {'color': '#FF0000', 'style': '-', 'width': 1.0, 'marker': None}
        }
        
        # 初始化界面布局
        self.init_ui()
        
        # 预填一些演示数据
        self.init_points()

    def init_ui(self):
        self.set_app_icon() # 设置程序图标
        # --- 菜单栏 ---
        menubar = self.menuBar()
        settings_menu = menubar.addMenu("设置")
        self.act_style = settings_menu.addAction("图表样式设置")
        self.act_style.triggered.connect(self.open_style_config)
        self.act_theme = settings_menu.addAction("切换暗黑模式")
        self.act_theme.triggered.connect(self.toggle_theme)
        self.act_api = settings_menu.addAction("API 设置")
        self.act_api.triggered.connect(self.open_api_settings)
        help_menu = menubar.addMenu("帮助")
        about_act = help_menu.addAction("关于")
        about_act.triggered.connect(self.show_about)
        self.lbl_github = QLabel(
            '本项目已在GitHub开源 <a href="https://github.com/ForestSun2023/SettlementPro">'
            'https://github.com/ForestSun2023/SettlementPro</a>'
        )
        self.lbl_github.setTextFormat(Qt.TextFormat.RichText)
        self.lbl_github.setTextInteractionFlags(Qt.TextInteractionFlag.TextBrowserInteraction)
        self.lbl_github.setOpenExternalLinks(True)
        self.lbl_github.setStyleSheet("margin-right: 50px; margin-top: 10px;")
        menubar.setCornerWidget(self.lbl_github, Qt.Corner.TopRightCorner)

        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        main_layout = QHBoxLayout(main_widget)

        # --- 左侧面板：数据输入与控制 ---
        left_panel = QVBoxLayout()
        
        # 1. 数据表格区域
        data_group = CollapsibleBox("监测数据录入")
        data_layout = QVBoxLayout()
        
        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["期数 (t)", "累积沉降量 (mm)", "施工阶段"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        self.table.setMinimumHeight(220)
        self.table.itemChanged.connect(self.on_table_changed)
        data_layout.addWidget(self.table)
        
        # 表格操作按钮
        btn_layout = QHBoxLayout()
        self.btn_add_row = QPushButton("添加行")
        self.btn_add_row.clicked.connect(self.add_row)
        self.btn_del_row = QPushButton("删除行")
        self.btn_del_row.clicked.connect(self.del_row)
        self.btn_clear = QPushButton("清空数据")
        self.btn_clear.clicked.connect(self.clear_data)
        self.btn_import = QPushButton("导入 CSV")
        self.btn_import.clicked.connect(self.import_csv)
        
        self.btn_save = QPushButton("保存数据")
        self.btn_save.clicked.connect(self.save_data)
        
        btn_layout.addWidget(self.btn_add_row)
        btn_layout.addWidget(self.btn_del_row)
        btn_layout.addWidget(self.btn_clear)
        btn_layout.addWidget(self.btn_import)
        btn_layout.addWidget(self.btn_save)
        data_layout.addLayout(btn_layout)
        data_group.setContentLayout(data_layout)
        
        # 1.5 测点管理区域
        point_group = CollapsibleBox("测点管理")
        point_layout = QVBoxLayout()
        
        point_select_layout = QHBoxLayout()
        point_select_layout.addWidget(QLabel("当前测点:"))
        self.combo_point = QComboBox()
        self.combo_point.currentTextChanged.connect(self.on_point_changed)
        point_select_layout.addWidget(self.combo_point)
        point_layout.addLayout(point_select_layout)
        
        point_btn_layout = QHBoxLayout()
        self.btn_add_point = QPushButton("新增测点")
        self.btn_add_point.clicked.connect(self.add_point)
        self.btn_rename_point = QPushButton("重命名")
        self.btn_rename_point.clicked.connect(self.rename_point)
        self.btn_del_point = QPushButton("删除测点")
        self.btn_del_point.clicked.connect(self.delete_point)
        point_btn_layout.addWidget(self.btn_add_point)
        point_btn_layout.addWidget(self.btn_rename_point)
        point_btn_layout.addWidget(self.btn_del_point)
        point_layout.addLayout(point_btn_layout)

        base_btn_layout = QHBoxLayout()
        self.btn_multi_ref = QPushButton("多基准设置")
        self.btn_multi_ref.clicked.connect(self.open_multi_ref_dialog)
        self.chk_apply_ref = QCheckBox("启用基准漂移修正")
        base_btn_layout.addWidget(self.btn_multi_ref)
        base_btn_layout.addWidget(self.chk_apply_ref)
        point_layout.addLayout(base_btn_layout)
        self.lbl_multi_ref = QLabel("多基准：未选择")
        point_layout.addWidget(self.lbl_multi_ref)
        
        pos_layout = QHBoxLayout()
        pos_layout.addWidget(QLabel("里程/位置 (m):"))
        self.input_point_pos = QDoubleSpinBox()
        self.input_point_pos.setRange(-100000.0, 100000.0)
        self.input_point_pos.setDecimals(3)
        self.input_point_pos.setValue(0.0)
        self.input_point_pos.valueChanged.connect(self.update_point_position)
        pos_layout.addWidget(self.input_point_pos)
        point_layout.addLayout(pos_layout)
        
        ref_layout = QHBoxLayout()
        ref_layout.addWidget(QLabel("参考点:"))
        self.combo_ref_point = QComboBox()
        self.combo_ref_point.currentTextChanged.connect(self.update_multi_point_metrics)
        ref_layout.addWidget(self.combo_ref_point)
        point_layout.addLayout(ref_layout)
        
        self.chk_show_ref = QCheckBox("显示参考点曲线")
        self.chk_show_ref.setChecked(True)
        point_layout.addWidget(self.chk_show_ref)
        
        point_group.setContentLayout(point_layout)
        
        # 2. 预测控制区域
        control_group = CollapsibleBox("模型计算与预测")
        control_layout = QVBoxLayout()
        
        h_layout = QHBoxLayout()
        h_layout.addWidget(QLabel("预测未来期数:"))
        self.input_future = QLineEdit("3")
        h_layout.addWidget(self.input_future)
        control_layout.addLayout(h_layout)

        model_layout = QHBoxLayout()
        model_layout.addWidget(QLabel("主模型:"))
        self.combo_model = QComboBox()
        self.combo_model.addItems(["自动选型", "GM(1,1)", "对数模型", "幂函数模型", "指数模型", "双曲线模型"])
        self.combo_model.setCurrentText("自动选型")
        model_layout.addWidget(self.combo_model)
        control_layout.addLayout(model_layout)

        self.lbl_model_select = QLabel("自动选型结果：--")
        self.lbl_model_rmse = QLabel("RMSE：--")
        self.lbl_model_mae = QLabel("MAE：--")
        self.lbl_model_note = QLabel("参数说明：--")
        control_layout.addWidget(self.lbl_model_select)
        control_layout.addWidget(self.lbl_model_rmse)
        control_layout.addWidget(self.lbl_model_mae)
        control_layout.addWidget(self.lbl_model_note)
        
        h_layout_thresh = QHBoxLayout()
        h_layout_thresh.addWidget(QLabel("速率阈值 (mm):"))
        self.input_rate_threshold = QLineEdit("2.0")
        h_layout_thresh.addWidget(self.input_rate_threshold)
        control_layout.addLayout(h_layout_thresh)

        self.chk_rolling = QCheckBox("启用滚动预测检验")
        control_layout.addWidget(self.chk_rolling)
        
        self.chk_kalman = QCheckBox("显示卡尔曼滤波 (去噪)")
        control_layout.addWidget(self.chk_kalman)

        # -- 修改：将移动平均复选框与窗口大小输入框组合 --
        ma_layout = QHBoxLayout()
        self.chk_moving_avg = QCheckBox("移动平均 (窗口):")
        self.input_ma_window = QLineEdit("3")
        self.input_ma_window.setFixedWidth(60) # 设置一个合适的宽度
        ma_layout.addWidget(self.chk_moving_avg)
        ma_layout.addWidget(self.input_ma_window)
        control_layout.addLayout(ma_layout)
        
        self.chk_ema = QCheckBox("显示指数平滑 (EMA)")
        control_layout.addWidget(self.chk_ema)
        
        # -- 新增：多项式拟合阶数选择 --
        poly_layout = QHBoxLayout()
        self.chk_poly = QCheckBox("显示多项式拟合 (阶数):")
        self.chk_poly.setChecked(True) # 默认开启
        self.combo_poly_order = QComboBox()
        self.combo_poly_order.addItems(["1", "2", "3", "4", "5"])
        self.combo_poly_order.setCurrentText("2") # 默认 2 阶
        self.combo_poly_order.setFixedWidth(60)
        poly_layout.addWidget(self.chk_poly)
        poly_layout.addWidget(self.combo_poly_order)
        control_layout.addLayout(poly_layout)
        
        self.chk_conf_int = QCheckBox("显示 95% 置信区间")
        self.chk_conf_int.setChecked(True)
        control_layout.addWidget(self.chk_conf_int)
        
        # -- 新增：异常值检测 --
        self.chk_detect_outliers = QCheckBox("启用异常值检测")
        self.chk_detect_outliers.setChecked(True)
        control_layout.addWidget(self.chk_detect_outliers)
        
        outlier_layout = QHBoxLayout()
        outlier_layout.addWidget(QLabel("检测方法:"))
        self.combo_outlier_method = QComboBox()
        self.combo_outlier_method.addItems(["IQR", "3σ", "Grubbs"])
        outlier_layout.addWidget(self.combo_outlier_method)
        outlier_layout.addWidget(QLabel("α:"))
        self.combo_grubbs_alpha = QComboBox()
        self.combo_grubbs_alpha.addItems(["0.05", "0.01"])
        self.combo_grubbs_alpha.setCurrentText("0.05")
        outlier_layout.addWidget(self.combo_grubbs_alpha)
        control_layout.addLayout(outlier_layout)

        outlier_handle_layout = QHBoxLayout()
        outlier_handle_layout.addWidget(QLabel("异常处理:"))
        self.combo_outlier_handle = QComboBox()
        self.combo_outlier_handle.addItems(["不处理", "剔除", "线性插值", "邻点均值"])
        outlier_handle_layout.addWidget(self.combo_outlier_handle)
        control_layout.addLayout(outlier_handle_layout)

        change_layout = QHBoxLayout()
        self.chk_change_detect = QCheckBox("工况识别 (CUSUM)")
        self.chk_change_detect.setChecked(True)
        change_layout.addWidget(self.chk_change_detect)
        change_layout.addWidget(QLabel("阈值:"))
        self.input_cusum_threshold = QLineEdit("3.0")
        self.input_cusum_threshold.setFixedWidth(60)
        change_layout.addWidget(self.input_cusum_threshold)
        control_layout.addLayout(change_layout)

        # -- 新增：日期设置 --
        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel("起始日期:"))
        self.input_start_date = QDateEdit()
        self.input_start_date.setDate(QDate.currentDate()) # 默认为今天
        self.input_start_date.setCalendarPopup(True) # 启用日历弹窗
        date_layout.addWidget(self.input_start_date)
        date_layout.addWidget(QLabel("间隔(天):"))
        self.input_interval = QLineEdit("7") # 默认7天一期
        self.input_interval.setFixedWidth(40)
        date_layout.addWidget(self.input_interval)
        control_layout.addLayout(date_layout)

        self.btn_calculate = QPushButton("执行计算与预测")
        self.btn_calculate.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 8px;")
        self.btn_calculate.clicked.connect(self.run_prediction)
        control_layout.addWidget(self.btn_calculate)
        
        self.btn_export_excel = QPushButton("导出 Excel")
        self.btn_export_excel.clicked.connect(self.export_to_excel)
        
        export_layout = QHBoxLayout()
        export_layout.addWidget(self.btn_export_excel)
        
        self.btn_export_word = QPushButton("生成 Word 报告")
        self.btn_export_word.clicked.connect(self.export_word_report)
        export_layout.addWidget(self.btn_export_word)

        self.btn_export = QPushButton("导出分析图表")
        self.btn_export.clicked.connect(self.export_image)
        export_layout.addWidget(self.btn_export)
        control_layout.addLayout(export_layout)
        
        control_group.setContentLayout(control_layout)
        
        # 2.5 多点评价参数
        eval_group = CollapsibleBox("多点评价参数")
        eval_layout = QGridLayout()
        
        eval_layout.addWidget(QLabel("稳定判定窗口 N:"), 0, 0)
        self.input_stable_window = QLineEdit("3")
        self.input_stable_window.setFixedWidth(60)
        eval_layout.addWidget(self.input_stable_window, 0, 1)
        
        eval_layout.addWidget(QLabel("稳定速率阈值 (mm/期):"), 0, 2)
        self.input_stable_threshold = QLineEdit("0.5")
        self.input_stable_threshold.setFixedWidth(60)
        eval_layout.addWidget(self.input_stable_threshold, 0, 3)
        
        eval_layout.addWidget(QLabel("累计沉降限值 (mm):"), 1, 0)
        self.input_total_limit = QLineEdit("30")
        self.input_total_limit.setFixedWidth(60)
        eval_layout.addWidget(self.input_total_limit, 1, 1)
        
        eval_layout.addWidget(QLabel("差异沉降限值 (mm):"), 1, 2)
        self.input_diff_limit = QLineEdit("15")
        self.input_diff_limit.setFixedWidth(60)
        eval_layout.addWidget(self.input_diff_limit, 1, 3)
        
        eval_layout.addWidget(QLabel("倾斜率限值 (mm/m):"), 2, 0)
        self.input_tilt_limit = QLineEdit("2.0")
        self.input_tilt_limit.setFixedWidth(60)
        eval_layout.addWidget(self.input_tilt_limit, 2, 1)

        eval_layout.addWidget(QLabel("基准漂移阈值 (mm):"), 2, 2)
        self.input_ref_drift_limit = QLineEdit("2.0")
        self.input_ref_drift_limit.setFixedWidth(60)
        eval_layout.addWidget(self.input_ref_drift_limit, 2, 3)

        eval_layout.addWidget(QLabel("角变形限值 (1/):"), 3, 0)
        self.input_ang_limit = QLineEdit("500")
        self.input_ang_limit.setFixedWidth(60)
        eval_layout.addWidget(self.input_ang_limit, 3, 1)
        
        eval_group.setContentLayout(eval_layout)

        # 2.55 观测精度评定
        acc_group = CollapsibleBox("观测精度评定")
        acc_layout = QGridLayout()
        acc_layout.addWidget(QLabel("仪器精度 (mm):"), 0, 0)
        self.input_inst_precision = QLineEdit("1.0")
        self.input_inst_precision.setFixedWidth(60)
        acc_layout.addWidget(self.input_inst_precision, 0, 1)
        acc_layout.addWidget(QLabel("测回数:"), 0, 2)
        self.input_rounds = QLineEdit("2")
        self.input_rounds.setFixedWidth(60)
        acc_layout.addWidget(self.input_rounds, 0, 3)
        acc_layout.addWidget(QLabel("观测标准差 (mm):"), 1, 0)
        self.input_obs_std = QLineEdit("1.0")
        self.input_obs_std.setFixedWidth(60)
        acc_layout.addWidget(self.input_obs_std, 1, 1)
        self.lbl_acc_eval = QLabel("评定：--")
        acc_layout.addWidget(self.lbl_acc_eval, 1, 2, 1, 2)
        acc_group.setContentLayout(acc_layout)

        # 2.6 变形指标展示
        metric_group = CollapsibleBox("多点变形指标")
        metric_layout = QVBoxLayout()
        self.lbl_diff_settlement = QLabel("差异沉降：-- mm")
        self.lbl_tilt = QLabel("倾斜率：-- mm/m")
        self.lbl_ang_distortion = QLabel("角变形：--")
        self.lbl_grade_diff = QLabel("差异沉降等级：--")
        self.lbl_grade_tilt = QLabel("倾斜等级：--")
        self.lbl_grade_ang = QLabel("角变形等级：--")
        self.lbl_stability = QLabel("稳定性：--")
        self.lbl_risk = QLabel("风险等级：--")
        metric_layout.addWidget(self.lbl_diff_settlement)
        metric_layout.addWidget(self.lbl_tilt)
        metric_layout.addWidget(self.lbl_ang_distortion)
        metric_layout.addWidget(self.lbl_grade_diff)
        metric_layout.addWidget(self.lbl_grade_tilt)
        metric_layout.addWidget(self.lbl_grade_ang)
        metric_layout.addWidget(self.lbl_stability)
        metric_layout.addWidget(self.lbl_risk)
        metric_group.setContentLayout(metric_layout)

        # 2.65 复测频率建议
        resurvey_group = CollapsibleBox("复测频率建议")
        resurvey_layout = QVBoxLayout()
        self.table_resurvey = QTableWidget()
        self.table_resurvey.setColumnCount(3)
        self.table_resurvey.setHorizontalHeaderLabels(["测点", "最近间距(m)", "建议频率"])
        self.table_resurvey.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table_resurvey.setMinimumHeight(120)
        resurvey_layout.addWidget(self.table_resurvey)
        resurvey_group.setContentLayout(resurvey_layout)

        # 2.7 协同分析
        coop_group = CollapsibleBox("多测点协同分析")
        coop_layout = QVBoxLayout()
        self.lbl_mean_curve = QLabel("群点平均曲线：--")
        self.lbl_max_diff = QLabel("最大差异点：--")
        self.lbl_outlier_spread = QLabel("异常传播：--")
        coop_layout.addWidget(self.lbl_mean_curve)
        coop_layout.addWidget(self.lbl_max_diff)
        coop_layout.addWidget(self.lbl_outlier_spread)
        coop_group.setContentLayout(coop_layout)

        # 3. 结果日志区域
        log_group = CollapsibleBox("计算报告")
        log_layout = QVBoxLayout()
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        self.log_text.setMinimumHeight(240)
        log_layout.addWidget(self.log_text)
        log_group.setContentLayout(log_layout)

        # 添加到左侧面板
        left_panel.addWidget(data_group, stretch=4)
        left_panel.addWidget(point_group, stretch=1)
        left_panel.addWidget(control_group, stretch=1)
        left_panel.addWidget(eval_group, stretch=1)
        left_panel.addWidget(metric_group, stretch=1)
        left_panel.addWidget(acc_group, stretch=1)
        left_panel.addWidget(resurvey_group, stretch=1)
        left_panel.addWidget(coop_group, stretch=1)
        left_panel.addWidget(log_group, stretch=3)

        # --- 右侧面板：绘图区域 ---
        right_panel = QVBoxLayout()
        plot_group = QGroupBox("沉降曲线可视化")
        plot_layout = QVBoxLayout()
        
        self.canvas = MplCanvas(self, width=5, height=4, dpi=100)
        plot_layout.addWidget(self.canvas)
        plot_group.setLayout(plot_layout)
        
        right_panel.addWidget(plot_group)

        # --- 组合主布局 ---
        left_container = QWidget()
        left_container.setLayout(left_panel)
        left_scroll = QScrollArea()
        left_scroll.setWidgetResizable(True)
        left_scroll.setWidget(left_container)
        left_scroll.setMinimumWidth(430)
        main_layout.addWidget(left_scroll, stretch=1)
        main_layout.addLayout(right_panel, stretch=2)

    def set_app_icon(self):
        """设置应用程序图标 (动态生成，无需外部文件)"""
        # 创建一个 64x64 的画布
        pixmap = QPixmap(64, 64)
        pixmap.fill(QColor(0, 0, 0, 0)) # 透明背景

        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        
        # 1. 绘制背景 (深青色圆角矩形)
        painter.setBrush(QColor("#00838F"))
        painter.setPen(Qt.PenStyle.NoPen)
        painter.drawRoundedRect(4, 4, 56, 56, 12, 12)
        
        # 2. 绘制折线图符号 (白色)
        pen = QPen(QColor("white"))
        pen.setWidth(5)
        pen.setCapStyle(Qt.PenCapStyle.RoundCap)
        pen.setJoinStyle(Qt.PenJoinStyle.RoundJoin)
        painter.setPen(pen)
        
        # 绘制简单的趋势线
        painter.drawLine(15, 48, 30, 32)
        painter.drawLine(30, 32, 42, 40)
        painter.drawLine(42, 40, 52, 16)
        
        painter.end()
        
        # 应用图标到窗口和任务栏
        icon = QIcon(pixmap)
        self.setWindowIcon(icon)
        QApplication.instance().setWindowIcon(icon)

    def toggle_theme(self):
        """切换暗黑/明亮模式"""
        app = QApplication.instance()
        if not self.is_dark_mode:
            # 切换到暗黑模式
            app.setStyleSheet(DARK_STYLESHEET)
            self.is_dark_mode = True
            if hasattr(self, "act_theme"):
                self.act_theme.setText("切换明亮模式")
        else:
            # 切换回默认模式
            app.setStyleSheet("")
            self.is_dark_mode = False
            if hasattr(self, "act_theme"):
                self.act_theme.setText("切换暗黑模式")
        
        # 刷新图表配色
        self.refresh_plot_theme()

    def open_style_config(self):
        """打开样式配置对话框"""
        dlg = StyleConfigDialog(self.plot_styles, self)
        if dlg.exec():
            self.plot_styles = dlg.get_styles()
            self.run_prediction() # 重新绘制以应用样式

    def refresh_plot_theme(self):
        """仅刷新图表颜色而不重绘数据"""
        bg_color = '#2b2b2b' if self.is_dark_mode else 'white'
        text_color = 'white' if self.is_dark_mode else 'black'
        
        self.canvas.figure.set_facecolor(bg_color)
        
        for ax in [self.canvas.ax1, self.canvas.ax2]:
            ax.set_facecolor(bg_color)
            ax.title.set_color(text_color)
            ax.xaxis.label.set_color(text_color)
            ax.yaxis.label.set_color(text_color)
            ax.tick_params(colors=text_color)
            for spine in ax.spines.values():
                spine.set_color(text_color)
            
            legend = ax.get_legend()
            if legend:
                for text in legend.get_texts(): text.set_color(text_color)
                legend.get_frame().set_facecolor('#3c3f41' if self.is_dark_mode else 'white')
        self.canvas.draw()

    def show_about(self):
        """显示关于对话框"""
        msg = QMessageBox(self)
        msg.setWindowTitle("关于")
        msg.setTextFormat(Qt.TextFormat.RichText)
        msg.setText(
            "<b>沉降监测与变形预测系统 v1.0</b><br><br>"
            "基于 Python PyQt6、灰色预测、曲线拟合、滚动预测与异常检测等技术开发。<br><br>"
            "本项目已在 GitHub 开源："
            "<a href=\"https://github.com/ForestSun2023/SettlementPro\">"
            "https://github.com/ForestSun2023/SettlementPro</a>"
        )
        msg.setTextInteractionFlags(Qt.TextInteractionFlag.TextBrowserInteraction)
        msg.setStandardButtons(QMessageBox.StandardButton.Ok)
        for label in msg.findChildren(QLabel):
            label.setOpenExternalLinks(True)
        msg.exec()

    def get_settings_path(self):
        return os.path.join(os.getcwd(), "settlement_settings.json")

    def load_api_settings(self):
        path = self.get_settings_path()
        if not os.path.exists(path):
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self.api_endpoint = data.get("endpoint", self.api_endpoint)
            self.api_model = data.get("model", self.api_model)
            self.api_key = data.get("api_key", self.api_key)
        except Exception:
            pass

    def save_api_settings(self):
        path = self.get_settings_path()
        data = {
            "endpoint": self.api_endpoint,
            "model": self.api_model,
            "api_key": self.api_key
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            QMessageBox.warning(self, "设置保存失败", f"无法保存 API 设置:\n{e}")

    def open_api_settings(self):
        settings = {"endpoint": self.api_endpoint, "model": self.api_model, "api_key": self.api_key}
        dlg = ApiSettingsDialog(settings, self)
        if dlg.exec():
            new_settings = dlg.get_settings()
            self.api_endpoint = new_settings.get("endpoint", self.api_endpoint)
            self.api_model = new_settings.get("model", self.api_model)
            self.api_key = new_settings.get("api_key", self.api_key)
            self.save_api_settings()

    def add_row(self):
        row_count = self.table.rowCount()
        self.table.insertRow(row_count)
        self.table.setItem(row_count, 0, QTableWidgetItem(str(row_count + 1)))
        self.table.setItem(row_count, 1, QTableWidgetItem("0.0"))
        self.table.setItem(row_count, 2, QTableWidgetItem(""))

    def del_row(self):
        current_row = self.table.currentRow()
        if current_row > -1:
            self.table.removeRow(current_row)
        else:
            self.table.removeRow(self.table.rowCount() - 1)
        self.save_current_point_data()
        self.update_multi_point_metrics()

    def clear_data(self):
        self.table.setRowCount(0)
        self.log_text.clear()
        self.canvas.ax1.cla()
        self.canvas.ax2.cla()
        self.canvas.draw()
        if self.current_point:
            self.points_data[self.current_point] = []
            self.points_stage[self.current_point] = []
        self.update_multi_point_metrics()

    def init_points(self):
        """初始化测点管理与演示数据"""
        self.points_data = {}
        self.points_stage = {}
        self.point_positions = {}
        for i in range(1, 6):
            name = f"P{i}"
            self.points_data[name] = []
            self.points_stage[name] = []
            self.point_positions[name] = float((i - 1) * 5)
        self.current_point = "P1"
        self.refresh_point_selectors()
        self.load_demo_data()

    def load_demo_data(self):
        """加载演示数据"""
        demo_sets = {
            "P1": ([10.5, 14.2, 17.8, 20.9, 23.6, 25.8], ["开挖", "开挖", "支护", "停工", "回填", "回填"]),
            "P2": ([9.8, 13.1, 16.4, 19.2, 22.0, 23.9], ["开挖", "开挖", "支护", "停工", "回填", "回填"]),
            "P3": ([11.2, 15.0, 18.0, 21.7, 24.5, 26.2], ["开挖", "开挖", "支护", "停工", "回填", "回填"]),
            "P4": ([8.9, 12.0, 15.5, 18.6, 20.7, 22.4], ["开挖", "开挖", "支护", "停工", "回填", "回填"]),
            "P5": ([10.0, 13.6, 16.9, 19.9, 22.8, 24.6], ["开挖", "开挖", "支护", "停工", "回填", "回填"]),
        }
        for name, (vals, stages) in demo_sets.items():
            self.points_data[name] = list(vals)
            self.points_stage[name] = list(stages)

        self.load_point_data(self.current_point)

    def import_csv(self):
        """从 CSV 或文本文件导入数据"""
        fp, _ = QFileDialog.getOpenFileName(self, "导入数据", "", "Data Files (*.csv *.txt);;All Files (*)")
        if not fp:
            return
            
        try:
            imported_data = []
            with open(fp, 'r', encoding='utf-8') as f:
                for line in f:
                    parts = [p.strip() for p in line.strip().split(',')]
                    if len(parts) >= 2:
                        try:
                            t_val = float(parts[0])
                            s_val = float(parts[1])
                            stage = parts[2] if len(parts) >= 3 else ""
                            imported_data.append((t_val, s_val, stage))
                            continue
                        except ValueError:
                            pass

                    # 兼容空格分隔或仅数值格式
                    parts = line.replace(',', ' ').split()
                    nums = []
                    for p in parts:
                        try: nums.append(float(p))
                        except ValueError: pass
                    
                    if len(nums) >= 2:
                        imported_data.append((nums[0], nums[1], "")) # 假设格式: 期数, 沉降量
                    elif len(nums) == 1:
                        imported_data.append((None, nums[0], ""))    # 假设格式: 仅沉降量
            
            if not imported_data:
                QMessageBox.warning(self, "导入失败", "未在文件中找到有效数据。\n请确保文件包含数字列。")
                return

            self.clear_data()
            for i, (t, s, stage) in enumerate(imported_data):
                self.add_row()
                # 如果文件包含期数则使用，否则自动生成 1, 2, 3...
                t_str = str(int(t)) if t is not None else str(i + 1)
                self.table.setItem(i, 0, QTableWidgetItem(t_str))
                self.table.setItem(i, 1, QTableWidgetItem(str(s)))
                self.table.setItem(i, 2, QTableWidgetItem(stage))
            
            self.save_current_point_data()
            QMessageBox.information(self, "导入成功", f"已加载 {len(imported_data)} 条记录。")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"读取文件时发生错误:\n{e}")

    def save_data(self):
        """导出表格数据到 CSV"""
        if self.table.rowCount() == 0:
            QMessageBox.warning(self, "提示", "表格为空，无数据可保存。")
            return

        fp, _ = QFileDialog.getSaveFileName(self, "保存数据", "settlement_data.csv", "CSV Files (*.csv);;All Files (*)")
        if not fp:
            return

        try:
            with open(fp, 'w', encoding='utf-8') as f:
                for row in range(self.table.rowCount()):
                    t_item = self.table.item(row, 0)
                    s_item = self.table.item(row, 1)
                    stage_item = self.table.item(row, 2)
                    if t_item and s_item:
                        stage_text = stage_item.text() if stage_item else ""
                        f.write(f"{t_item.text()},{s_item.text()},{stage_text}\n")
            QMessageBox.information(self, "成功", f"数据已保存至:\n{fp}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存文件时发生错误:\n{e}")

    def get_table_data(self, allow_empty=False):
        """从表格获取数据"""
        data = []
        try:
            for row in range(self.table.rowCount()):
                item = self.table.item(row, 1)
                if item and item.text():
                    data.append(float(item.text()))
            if not data and not allow_empty:
                return None
            return data
        except ValueError:
            QMessageBox.warning(self, "输入错误", "请确保沉降量为有效的数字！")
            return None

    def get_stage_labels(self):
        labels = []
        for row in range(self.table.rowCount()):
            item = self.table.item(row, 2)
            label = item.text().strip() if item and item.text() else "未标注"
            labels.append(label)
        return labels

    def open_multi_ref_dialog(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("多基准设置")
        layout = QVBoxLayout(dlg)
        layout.addWidget(QLabel("请选择用于联合平差的基准点 (至少2个)："))
        checks = []
        for name in self.points_data.keys():
            cb = QCheckBox(name)
            cb.setChecked(name in self.multi_ref_points)
            layout.addWidget(cb)
            checks.append(cb)

        btn_layout = QHBoxLayout()
        btn_ok = QPushButton("确定")
        btn_cancel = QPushButton("取消")
        btn_ok.clicked.connect(dlg.accept)
        btn_cancel.clicked.connect(dlg.reject)
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        if dlg.exec():
            selected = [c.text() for c in checks if c.isChecked()]
            self.multi_ref_points = selected
            if selected:
                self.lbl_multi_ref.setText(f"多基准：{', '.join(selected)}")
            else:
                self.lbl_multi_ref.setText("多基准：未选择")

    def save_current_point_data(self):
        """保存当前表格数据到测点"""
        if not self.current_point or self.is_loading_point:
            return
        data = self.get_table_data(allow_empty=True)
        if data is None:
            return
        self.points_data[self.current_point] = data
        self.points_stage[self.current_point] = self.get_stage_labels()

    def on_table_changed(self, item):
        if self.is_loading_point:
            return
        self.save_current_point_data()
        self.update_multi_point_metrics()

    def load_point_data(self, point_name):
        """加载指定测点数据到表格"""
        if point_name not in self.points_data:
            return
        self.is_loading_point = True
        self.table.setRowCount(0)
        data = self.points_data.get(point_name, [])
        stage_labels = self.points_stage.get(point_name, [])
        for i, val in enumerate(data):
            self.add_row()
            self.table.setItem(i, 0, QTableWidgetItem(str(i + 1)))
            self.table.setItem(i, 1, QTableWidgetItem(str(val)))
            stage_text = stage_labels[i] if i < len(stage_labels) else ""
            self.table.setItem(i, 2, QTableWidgetItem(stage_text))
        self.current_point = point_name
        self.input_point_pos.setValue(self.point_positions.get(point_name, 0.0))
        self.is_loading_point = False
        self.update_multi_point_metrics()

    def refresh_point_selectors(self):
        """刷新测点与参考点下拉框"""
        self.combo_point.blockSignals(True)
        self.combo_ref_point.blockSignals(True)
        self.combo_point.clear()
        self.combo_ref_point.clear()
        for name in self.points_data.keys():
            self.combo_point.addItem(name)
            self.combo_ref_point.addItem(name)
        if self.current_point:
            self.combo_point.setCurrentText(self.current_point)
        if self.current_point and self.current_point in self.points_data:
            self.combo_ref_point.setCurrentText(self.current_point)
        self.combo_point.blockSignals(False)
        self.combo_ref_point.blockSignals(False)

    def on_point_changed(self, point_name):
        if not point_name or point_name == self.current_point:
            return
        self.save_current_point_data()
        self.load_point_data(point_name)
        if self.point_results.get(point_name):
            self.display_point_results(point_name)

    def add_point(self):
        name, ok = QInputDialog.getText(self, "新增测点", "输入测点名称:")
        if not ok or not name.strip():
            return
        name = name.strip()
        if name in self.points_data:
            QMessageBox.warning(self, "重复测点", "测点名称已存在，请更换名称。")
            return
        self.save_current_point_data()
        self.points_data[name] = []
        self.points_stage[name] = []
        self.point_positions[name] = 0.0
        self.current_point = name
        self.refresh_point_selectors()
        self.load_point_data(name)

    def rename_point(self):
        if not self.current_point:
            return
        new_name, ok = QInputDialog.getText(self, "重命名测点", "输入新的测点名称:", text=self.current_point)
        if not ok or not new_name.strip():
            return
        new_name = new_name.strip()
        if new_name in self.points_data and new_name != self.current_point:
            QMessageBox.warning(self, "重复测点", "测点名称已存在，请更换名称。")
            return
        if new_name == self.current_point:
            return
        self.points_data[new_name] = self.points_data.pop(self.current_point)
        self.points_stage[new_name] = self.points_stage.pop(self.current_point, [])
        self.point_positions[new_name] = self.point_positions.pop(self.current_point, 0.0)
        self.current_point = new_name
        self.refresh_point_selectors()

    def delete_point(self):
        if not self.current_point:
            return
        if len(self.points_data) <= 1:
            QMessageBox.warning(self, "操作受限", "至少需要保留一个测点。")
            return
        reply = QMessageBox.question(self, "删除测点", f"确定删除测点 {self.current_point} ?")
        if reply != QMessageBox.StandardButton.Yes:
            return
        self.points_data.pop(self.current_point, None)
        self.points_stage.pop(self.current_point, None)
        self.point_positions.pop(self.current_point, None)
        self.current_point = next(iter(self.points_data.keys()))
        self.refresh_point_selectors()
        self.load_point_data(self.current_point)

    def update_point_position(self):
        if not self.current_point:
            return
        self.point_positions[self.current_point] = float(self.input_point_pos.value())
        self.update_multi_point_metrics()

    def detect_outliers(self, data, method="IQR", alpha=0.05):
        """
        基于速率的异常检测，支持 IQR/3σ/Grubbs。
        返回异常值点的索引列表。
        """
        if len(data) < 4:
            return []

        diffs = np.diff(data)
        method = method.upper()

        if method == "3Σ" or method == "3σ" or method == "3SIGMA":
            mean = np.mean(diffs)
            std = np.std(diffs)
            if std == 0:
                return []
            outlier_indices = np.where(np.abs(diffs - mean) > 3 * std)[0]
            return (outlier_indices + 1).tolist()

        if method == "GRUBBS":
            n = len(diffs)
            if n < 3:
                return []
            mean = np.mean(diffs)
            std = np.std(diffs)
            if std == 0:
                return []
            g_vals = np.abs(diffs - mean) / std
            max_idx = int(np.argmax(g_vals))
            g_max = g_vals[max_idx]
            g_crit = self.get_grubbs_critical_value(n, alpha=alpha)
            if g_crit is not None and g_max > g_crit:
                return [max_idx + 1]
            return []

        # 默认 IQR
        q1 = np.percentile(diffs, 25)
        q3 = np.percentile(diffs, 75)
        iqr = q3 - q1
        lower_bound = q1 - 1.5 * iqr
        upper_bound = q3 + 1.5 * iqr
        outlier_indices = np.where((diffs < lower_bound) | (diffs > upper_bound))[0]
        return (outlier_indices + 1).tolist()

    def get_grubbs_critical_value(self, n, alpha=0.05):
        # 近似 Grubbs 临界值，n=4..20
        table_005 = {
            4: 1.463, 5: 1.672, 6: 1.822, 7: 1.938, 8: 2.032, 9: 2.110, 10: 2.176,
            11: 2.234, 12: 2.285, 13: 2.331, 14: 2.371, 15: 2.409, 16: 2.443,
            17: 2.475, 18: 2.504, 19: 2.532, 20: 2.557
        }
        table_001 = {
            4: 1.672, 5: 1.917, 6: 2.078, 7: 2.207, 8: 2.314, 9: 2.404, 10: 2.482,
            11: 2.550, 12: 2.610, 13: 2.663, 14: 2.710, 15: 2.752, 16: 2.790,
            17: 2.824, 18: 2.855, 19: 2.883, 20: 2.908
        }
        table = table_001 if alpha <= 0.01 else table_005
        if n in table:
            return table[n]
        if n > 20:
            return 2.9 if alpha <= 0.01 else 2.6
        return None

    def highlight_outliers(self, outlier_indices):
        default_bg = QColor("#3c3f41") if self.is_dark_mode else QColor("white")
        for row in range(self.table.rowCount()):
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                if item:
                    item.setBackground(default_bg)
        for idx in outlier_indices:
            if 0 <= idx < self.table.rowCount():
                for col in range(self.table.columnCount()):
                    item = self.table.item(idx, col)
                    if item:
                        item.setBackground(QColor("#FFCDD2"))

    def detect_change_point_cusum(self, data, threshold):
        if len(data) < 4:
            return None
        diffs = np.diff(data)
        mean = np.mean(diffs)
        cusum = 0.0
        for i, d in enumerate(diffs):
            cusum += d - mean
            if abs(cusum) >= threshold:
                return i + 1
        return None

    def compute_baseline_series(self, length):
        if len(self.multi_ref_points) < 2:
            return None
        series_list = []
        for name in self.multi_ref_points:
            data = self.points_data.get(name, [])
            if len(data) >= length:
                series_list.append(np.array(data[:length], dtype=float))
        if len(series_list) < 2:
            return None
        return np.mean(series_list, axis=0)

    def compute_stage_stats(self, raw_data, stage_labels):
        if len(raw_data) < 2:
            return None, None
        diffs = np.diff(raw_data)
        stage_rates = {}
        stage_order = []
        for i in range(1, len(raw_data)):
            stage = stage_labels[i] if i < len(stage_labels) else "未标注"
            stage_rates.setdefault(stage, []).append(diffs[i - 1])
            if stage not in stage_order:
                stage_order.append(stage)

        stats = []
        for stage in stage_order:
            vals = np.array(stage_rates[stage], dtype=float)
            stats.append({
                "stage": stage,
                "count": len(vals),
                "avg_rate": float(np.mean(vals)),
                "max_rate": float(np.max(vals)),
                "min_rate": float(np.min(vals))
            })

        segments = []
        start = 0
        current = stage_labels[0] if stage_labels else "未标注"
        for i in range(1, len(stage_labels)):
            if stage_labels[i] != current:
                segments.append((current, start, i - 1))
                current = stage_labels[i]
                start = i
        segments.append((current, start, len(stage_labels) - 1))
        return stats, segments

    def compute_stage_deep_analysis(self, stage_stats):
        if not stage_stats or len(stage_stats) < 2:
            return None
        avg_abs = [abs(s["avg_rate"]) for s in stage_stats]
        overall = np.mean(avg_abs) if avg_abs else 1.0
        if overall == 0:
            overall = 1.0

        results = []
        prev_avg = None
        for s in stage_stats:
            avg_rate = s["avg_rate"]
            change_rate = None
            if prev_avg is not None and prev_avg != 0:
                change_rate = (avg_rate - prev_avg) / abs(prev_avg)
            impact = abs(avg_rate) / overall
            results.append({
                "stage": s["stage"],
                "avg_rate": avg_rate,
                "change_rate": change_rate,
                "impact": impact
            })
            prev_avg = avg_rate

        ranked = sorted(results, key=lambda x: x["impact"], reverse=True)
        rank_map = {r["stage"]: i + 1 for i, r in enumerate(ranked)}
        for r in results:
            r["rank"] = rank_map.get(r["stage"], "")
        return results

    def apply_outlier_strategy(self, raw_data, outlier_indices, strategy):
        if not outlier_indices or strategy == "不处理":
            return raw_data[:], False

        data = raw_data[:]
        outlier_set = set(outlier_indices)

        def get_prev_next(idx):
            prev_idx = idx - 1
            while prev_idx >= 0 and prev_idx in outlier_set:
                prev_idx -= 1
            next_idx = idx + 1
            while next_idx < len(data) and next_idx in outlier_set:
                next_idx += 1
            return prev_idx if prev_idx >= 0 else None, next_idx if next_idx < len(data) else None

        for idx in outlier_indices:
            prev_idx, next_idx = get_prev_next(idx)
            prev_val = data[prev_idx] if prev_idx is not None else None
            next_val = data[next_idx] if next_idx is not None else None
            if strategy in ("线性插值", "剔除"):
                if prev_val is not None and next_val is not None:
                    ratio = (idx - prev_idx) / (next_idx - prev_idx)
                    data[idx] = prev_val + (next_val - prev_val) * ratio
                elif prev_val is not None:
                    data[idx] = prev_val
                elif next_val is not None:
                    data[idx] = next_val
            elif strategy == "邻点均值":
                vals = [v for v in [prev_val, next_val] if v is not None]
                if vals:
                    data[idx] = float(np.mean(vals))
        return data, True

    def build_monitor_brief(self):
        items = []
        suggestions = []
        if self.last_outlier_indices is not None and len(self.last_outlier_indices) > 0:
            idx_str = ", ".join([str(i + 1) for i in list(self.last_outlier_indices)])
            items.append(f"异常点期次: {idx_str}")
            suggestions.append(f"建议复测: {self.current_point}")
        if self.last_risk_text:
            items.append(f"风险等级: {self.last_risk_text}")
        if self.last_stability_text:
            items.append(f"稳定性: {self.last_stability_text}")
        if self.last_unstable_refs:
            names = ", ".join(self.last_unstable_refs)
            items.append(f"基准稳定性: {names} 疑似漂移")
            suggestions.append(f"建议复测: {names}")
        if not items:
            items.append("异常与风险: 未发现明显问题")
        if not suggestions:
            suggestions.append("建议复测: 无")
        return items, suggestions

    def build_monitor_brief_all(self):
        items = []
        suggestions = []
        for name, res in self.point_results.items():
            if res.get("error"):
                continue
            out_cnt = len(res.get("outliers", []))
            risk = res.get("risk", "--")
            stability = res.get("stability", "--")
            items.append(f"{name}: 风险 {risk}, 稳定性 {stability}, 异常点 {out_cnt} 个")
            if out_cnt > 0 or risk in ("红色", "黄色") or stability == "未趋稳":
                suggestions.append(f"建议复测: {name}")
        if not items:
            items.append("异常与风险: 未发现明显问题")
        if not suggestions:
            suggestions.append("建议复测: 无")
        return items, suggestions

    def generate_ai_summary_all(self, brief_items, brief_suggestions):
        if not self.api_key:
            return None, "未配置 API Key"
        lines = []
        for name, res in self.point_results.items():
            if res.get("error"):
                continue
            raw_used = res.get("raw_used", [])
            last_val = raw_used[-1] if raw_used else 0.0
            line = f"{name}: 末期沉降 {last_val:.2f} mm, 模型 {res.get('model_name','')}, RMSE {res.get('rmse',0):.3f}, 风险 {res.get('risk','--')}, 稳定性 {res.get('stability','--')}"
            if res.get("change_index") is not None:
                line += f", 工况变化点 第{res['change_index']+1}期"
            lines.append(line)
            if res.get("stage_stats"):
                stage_info = "; ".join([f"{s['stage']} 平均速率 {s['avg_rate']:.3f}" for s in res["stage_stats"]])
                lines.append(f"{name} 施工阶段速率: {stage_info}")

        prompt = (
            "你是测绘工程监测报告助手。请基于以下要点生成更实际、更专业的中文总结，"
            "不少于500字，分为3段。内容需包含：趋势判断、阶段影响解读、异常与风险提示、"
            "复测与监测建议、工程影响与控制建议。避免空泛。\n"
            + "\n".join(lines)
            + "\n简报要点:\n"
            + "\n".join(brief_items + brief_suggestions)
        )

        payload = {
            "model": self.api_model,
            "messages": [
                {"role": "system", "content": "你是工程监测报告写作助手。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.2,
            "max_tokens": 1200
        }
        req_data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            self.api_endpoint,
            data=req_data,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            },
            method="POST"
        )
        timeout_sec = 60
        try:
            with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            content = data["choices"][0]["message"]["content"].strip()
            return content, None
        except urllib.error.HTTPError as e:
            body = ""
            try:
                body = e.read().decode("utf-8", errors="ignore")
            except Exception:
                body = ""
            msg = f"HTTP {e.code} | {e.reason} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"
            if body:
                msg += f" | body={body[:500]}"
            return None, msg
        except urllib.error.URLError as e:
            return None, f"URLError {e.reason} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"
        except Exception as e:
            return None, f"{e} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"

    def compute_optional_layers(self, raw_data, total_periods=None):
        poly_vals = None
        rolling_vals = None
        kalman_vals = None
        ma_vals = None
        ema_vals = None

        if self.chk_poly.isChecked():
            try:
                poly_order = int(self.combo_poly_order.currentText())
            except ValueError:
                poly_order = 2
            if len(raw_data) > poly_order:
                x = np.arange(1, len(raw_data) + 1)
                coeffs = np.polyfit(x, raw_data, poly_order)
                poly_fn = np.poly1d(coeffs)
                target_len = total_periods if total_periods else len(raw_data)
                poly_vals = poly_fn(np.arange(1, target_len + 1))

        if self.chk_rolling.isChecked() and len(raw_data) > 4:
            rolling_vals = self.gm_model.rolling_predict_check(raw_data)

        if self.chk_kalman.isChecked():
            kalman_filter = SimpleKalmanFilter()
            kalman_vals = kalman_filter.filter(raw_data)

        if self.chk_moving_avg.isChecked():
            try:
                window_size = int(self.input_ma_window.text())
            except ValueError:
                window_size = 3
            ma_filter = MovingAverageFilter(window_size=window_size)
            ma_vals = ma_filter.filter(raw_data)

        if self.chk_ema.isChecked():
            ema_filter = ExponentialMovingAverageFilter(alpha=0.4)
            ema_vals = ema_filter.filter(raw_data)

        return poly_vals, rolling_vals, kalman_vals, ma_vals, ema_vals

    def compute_point_result(self, point_name, raw_data_original, stage_labels, start_date, interval_days, future_periods, apply_ref, drift_limit, outlier_method, outlier_alpha, cusum_threshold, outlier_strategy):
        result = {"point": point_name}
        if not raw_data_original or len(raw_data_original) < 4:
            result["error"] = "数据不足"
            return result

        outlier_indices = []
        if self.chk_detect_outliers.isChecked():
            outlier_indices = self.detect_outliers(raw_data_original, method=outlier_method, alpha=outlier_alpha)

        raw_data = raw_data_original[:]
        raw_data, outlier_applied = self.apply_outlier_strategy(raw_data, outlier_indices, outlier_strategy)
        baseline = None
        if apply_ref:
            baseline = self.compute_baseline_series(len(raw_data_original))
            if baseline is not None:
                raw_data = (np.array(raw_data) - baseline).tolist()

        change_index = None
        if self.chk_change_detect.isChecked():
            change_index = self.detect_change_point_cusum(raw_data_original, cusum_threshold)

        total_periods = len(raw_data) + future_periods
        date_list = [start_date + timedelta(days=i * interval_days) for i in range(total_periods)]
        candidates = self.evaluate_models(raw_data, future_periods)
        if not candidates:
            result["error"] = "模型拟合失败"
            return result

        model_choice = self.combo_model.currentText()
        if model_choice == "自动选型":
            selected = min(candidates, key=lambda x: x["rmse"])
        else:
            selected = next((m for m in candidates if m["name"] == model_choice), None)
            if selected is None:
                selected = min(candidates, key=lambda x: x["rmse"])

        predicted_vals = selected["pred"]
        residuals = selected["residuals"]
        rel_errors = selected["rel_errors"]
        rmse = selected["rmse"]
        mae = float(np.mean(np.abs(residuals)))

        conf_bounds = None
        if self.chk_conf_int.isChecked():
            std_resid = np.std(residuals)
            margin = 1.96 * std_resid
            conf_bounds = (predicted_vals - margin, predicted_vals + margin)

        stage_stats, stage_segments = self.compute_stage_stats(raw_data_original, stage_labels)
        stage_deep = self.compute_stage_deep_analysis(stage_stats)
        metrics = self.compute_metrics_for_point(point_name, raw_data_original)
        risk, stability = self.assess_risk_stability(raw_data_original, metrics)
        grades = self.grade_deformation(metrics)
        try:
            stable_window = int(self.input_stable_window.text())
        except ValueError:
            stable_window = 3
        try:
            stable_threshold = float(self.input_stable_threshold.text())
        except ValueError:
            stable_threshold = 0.5
        stable_date = self.predict_stable_date(predicted_vals, date_list, stable_threshold, stable_window)

        result.update({
            "raw_original": raw_data_original,
            "raw_used": raw_data,
            "pred": predicted_vals,
            "residuals": residuals,
            "rel_errors": rel_errors,
            "rmse": rmse,
            "mae": mae,
            "model_name": selected["name"],
            "model_params": selected["params"],
            "conf_bounds": conf_bounds,
            "date_list": date_list,
            "outliers": outlier_indices,
            "change_index": change_index,
            "baseline": baseline,
            "corrected": apply_ref and baseline is not None,
            "stage_stats": stage_stats,
            "stage_segments": stage_segments,
            "stage_deep": stage_deep,
            "metrics": metrics,
            "risk": risk,
            "stability": stability,
            "grades": grades,
            "outlier_strategy": outlier_strategy,
            "outlier_applied": outlier_applied,
            "stable_date": stable_date
        })
        return result

    def compute_all_points(self):
        self.point_results = {}
        self.last_unstable_refs = []

        try:
            interval_days = int(self.input_interval.text())
        except ValueError:
            interval_days = 7
        q_date = self.input_start_date.date()
        start_date = datetime(q_date.year(), q_date.month(), q_date.day())
        try:
            future_periods = int(self.input_future.text())
        except ValueError:
            future_periods = 3
            self.input_future.setText("3")

        try:
            cusum_threshold = float(self.input_cusum_threshold.text())
        except ValueError:
            cusum_threshold = 3.0

        outlier_method = self.combo_outlier_method.currentText()
        try:
            outlier_alpha = float(self.combo_grubbs_alpha.currentText())
        except ValueError:
            outlier_alpha = 0.05
        outlier_strategy = self.combo_outlier_handle.currentText()

        apply_ref = self.chk_apply_ref.isChecked()
        try:
            drift_limit = float(self.input_ref_drift_limit.text())
        except ValueError:
            drift_limit = 2.0

        if apply_ref:
            ref_lengths = [len(self.points_data.get(n, [])) for n in self.multi_ref_points]
            if ref_lengths:
                length_ref = min(ref_lengths)
                unstable_refs = self.evaluate_reference_stability(length_ref, drift_limit)
                if unstable_refs:
                    self.last_unstable_refs = unstable_refs

        for name, data in self.points_data.items():
            stage_labels = self.points_stage.get(name, [])
            result = self.compute_point_result(
                name,
                data,
                stage_labels,
                start_date,
                interval_days,
                future_periods,
                apply_ref,
                drift_limit,
                outlier_method,
                outlier_alpha,
                cusum_threshold,
                outlier_strategy
            )
            self.point_results[name] = result

    def display_point_results(self, point_name):
        result = self.point_results.get(point_name)
        if not result or result.get("error"):
            QMessageBox.warning(self, "数据不足", f"{point_name}: {result.get('error', '无结果')}")
            return

        self.last_raw_data = result["raw_used"]
        self.last_predicted = result["pred"]
        self.last_residuals = result["residuals"]
        self.last_rel_errors = result["rel_errors"]
        self.last_model_name = result["model_name"]
        self.last_model_params = result["model_params"]
        self.last_model_rmse = result["rmse"]
        self.last_model_mae = result["mae"]
        self.last_change_index = result["change_index"]
        self.last_stage_stats = result["stage_stats"]
        self.last_stage_segments = result["stage_segments"]
        self.last_stage_deep = result.get("stage_deep")
        self.last_outlier_indices = result["outliers"]
        self.last_corrected = result["corrected"]
        self.last_outlier_applied = result.get("outlier_applied", False)
        self.last_risk_text = result.get("risk")
        self.last_stability_text = result.get("stability")
        self.last_deform_grades = result.get("grades")
        self.last_stable_date = result.get("stable_date")

        self.lbl_model_select.setText(f"自动选型结果：{self.last_model_name}")
        self.lbl_model_rmse.setText(f"RMSE：{self.last_model_rmse:.4f}")
        self.lbl_model_mae.setText(f"MAE：{self.last_model_mae:.4f}")

        note = "参数说明：--"
        if self.last_model_name == "指数模型" and self.last_model_params:
            b_val = self.last_model_params.get("b", 0.0)
            if b_val < 0:
                note = "参数说明：指数模型趋于 0（稳定值）"
            else:
                note = "参数说明：指数模型为发散趋势"
        elif self.last_model_name == "双曲线模型" and self.last_model_params:
            b_val = self.last_model_params.get("b", 0.0)
            if b_val != 0:
                stable = 1.0 / b_val
                note = f"参数说明：双曲线稳定值约为 {stable:.4f}"
        if self.chk_apply_ref.isChecked():
            note += " | 基准修正" if result["baseline"] is not None else " | 基准不足"
        self.lbl_model_note.setText(note)

        total_periods = len(result["pred"])
        poly_vals, rolling_vals, kalman_vals, ma_vals, ema_vals = self.compute_optional_layers(result["raw_used"], total_periods=total_periods)
        try:
            rate_threshold = float(self.input_rate_threshold.text())
        except ValueError:
            rate_threshold = 2.0
            self.input_rate_threshold.setText("2.0")
        try:
            interval_days = int(self.input_interval.text())
        except ValueError:
            interval_days = 7

        ref_name = self.combo_ref_point.currentText()
        ref_data = None
        if self.chk_show_ref.isChecked() and ref_name in self.points_data and ref_name != point_name:
            ref_data = self.points_data.get(ref_name, [])

        self.highlight_outliers(result["outliers"])
        self.update_plot(
            result["raw_used"],
            result["pred"],
            poly_vals,
            rolling_vals,
            kalman_vals,
            ma_vals,
            rate_threshold,
            ema_vals,
            result["conf_bounds"],
            result["date_list"],
            interval_days,
            result["outliers"],
            ref_data,
            ref_name,
            self.last_model_name,
            self.last_change_index,
            result["corrected"],
            result["raw_original"],
            result["baseline"],
            result["stage_stats"],
            result.get("outlier_applied", False),
            result.get("stable_date")
        )
        is_qualified, grade_msg = self.residual_check_generic(result["rel_errors"])
        is_qualified_cp, grade_msg_cp, c_val, p_val = self.posterior_variance_check_generic(result["raw_used"], result["residuals"])
        self.update_log(result["raw_used"], result["pred"], result["residuals"], result["rel_errors"],
                        grade_msg, grade_msg_cp, rate_threshold, result["date_list"])
        self.update_multi_point_metrics()
        self.update_accuracy_evaluation()
        self.update_resurvey_table()
        self.multi_point_analysis = self.compute_multi_point_analysis()
        self.update_multi_point_analysis_ui()

    def generate_ai_summary(self, raw, pred, metrics, stage_stats, brief_items, brief_suggestions):
        if not self.api_key:
            return None, "未配置 API Key"
        timeout_sec = 60
        prompt_lines = []
        if self.current_point:
            prompt_lines.append(f"当前测点: {self.current_point}")
        if self.last_risk_text:
            prompt_lines.append(f"风险等级: {self.last_risk_text}")
        if self.last_stability_text:
            prompt_lines.append(f"稳定性: {self.last_stability_text}")
        if self.last_outlier_indices is not None and len(self.last_outlier_indices) > 0:
            idx_str = ", ".join([str(i + 1) for i in list(self.last_outlier_indices)])
            prompt_lines.append(f"异常点期次: {idx_str}")
        if self.last_change_index is not None:
            prompt_lines.append(f"工况变化点: 第 {self.last_change_index + 1} 期")
        if metrics:
            prompt_lines.append(f"差异沉降: {metrics['diff']:+.2f} mm")
            if metrics.get("gradient") is not None:
                prompt_lines.append(f"倾斜率: {metrics['gradient']:.3f} mm/m")
        if stage_stats:
            stage_info = "; ".join([f"{s['stage']} 平均速率 {s['avg_rate']:.3f}" for s in stage_stats])
            prompt_lines.append(f"施工阶段速率: {stage_info}")

        prompt = (
            "你是测绘工程监测报告助手。请基于以下要点生成更实际、更专业的中文总结，"
            "不少于500字，分为3段。内容需包含：趋势判断、阶段影响解读、异常与风险提示、"
            "复测与监测建议、工程影响与控制建议。避免空泛。\n"
            + "\n".join(prompt_lines)
            + "\n简报要点:\n"
            + "\n".join(brief_items + brief_suggestions)
        )

        payload = {
            "model": self.api_model,
            "messages": [
                {"role": "system", "content": "你是工程监测报告写作助手。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.2,
            "max_tokens": 1200
        }
        req_data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            self.api_endpoint,
            data=req_data,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            },
            method="POST"
        )
        try:
            with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            content = data["choices"][0]["message"]["content"].strip()
            return content, None
        except urllib.error.HTTPError as e:
            body = ""
            try:
                body = e.read().decode("utf-8", errors="ignore")
            except Exception:
                body = ""
            msg = f"HTTP {e.code} | {e.reason} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"
            if body:
                msg += f" | body={body[:500]}"
            return None, msg
        except urllib.error.URLError as e:
            return None, f"URLError {e.reason} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"
        except Exception as e:
            return None, f"{e} | endpoint={self.api_endpoint} | model={self.api_model} | timeout={timeout_sec}s"

    def build_multi_point_figure(self):
        fig = plt.figure(figsize=(5.5, 3.2))
        ax = fig.add_subplot(111)
        for name, data in self.points_data.items():
            if not data:
                continue
            x = np.arange(1, len(data) + 1)
            if name == self.current_point:
                ax.plot(x, data, linewidth=2.0, label=name)
            else:
                ax.plot(x, data, linewidth=1.0, alpha=0.7, label=name)
        ax.set_title("多点沉降对比曲线")
        ax.set_xlabel("期数")
        ax.set_ylabel("沉降量 (mm)")
        ax.grid(True, linestyle="--", alpha=0.4)
        ax.legend(fontsize=7, ncol=3)
        fig.tight_layout()
        return fig

    def build_diff_tilt_figure(self, raw_data):
        ref_name = self.combo_ref_point.currentText()
        if not ref_name or ref_name == self.current_point:
            return None
        ref_data = self.points_data.get(ref_name, [])
        n = min(len(raw_data), len(ref_data))
        if n < 2:
            return None
        diff_series = np.array(raw_data[:n]) - np.array(ref_data[:n])
        distance = abs(self.point_positions.get(self.current_point, 0.0) - self.point_positions.get(ref_name, 0.0))
        if distance <= 0:
            return None
        tilt_series = diff_series / distance
        x = np.arange(1, n + 1)
        fig = plt.figure(figsize=(5.5, 3.2))
        ax = fig.add_subplot(111)
        ax.plot(x, diff_series, label="差异沉降 (mm)")
        ax.plot(x, tilt_series, label="倾斜率 (mm/m)")
        ax.set_title(f"差异沉降/倾斜趋势 ({self.current_point}-{ref_name})")
        ax.set_xlabel("期数")
        ax.grid(True, linestyle="--", alpha=0.4)
        ax.legend(fontsize=7)
        fig.tight_layout()
        return fig

    def build_outlier_compare_figure(self, raw_original, raw_used, point_name):
        if not raw_original or not raw_used:
            return None
        n = min(len(raw_original), len(raw_used))
        if n < 2:
            return None
        x = np.arange(1, n + 1)
        fig = plt.figure(figsize=(5.5, 3.2))
        ax = fig.add_subplot(111)
        ax.plot(x, raw_original[:n], label="处理前", linestyle="--", color="#9E9E9E")
        ax.plot(x, raw_used[:n], label="处理后", linestyle="-", color="#4A90E2")
        ax.set_title(f"异常点处理对比 ({point_name})")
        ax.set_xlabel("期数")
        ax.set_ylabel("沉降量 (mm)")
        ax.grid(True, linestyle="--", alpha=0.4)
        ax.legend(fontsize=7)
        fig.tight_layout()
        return fig

    def evaluate_reference_stability(self, length, drift_limit):
        if len(self.multi_ref_points) < 2:
            self.last_ref_stability_text = None
            return None
        base = self.compute_baseline_series(length)
        if base is None:
            self.last_ref_stability_text = None
            return None
        messages = []
        unstable = []
        for name in self.multi_ref_points:
            data = self.points_data.get(name, [])
            if len(data) < length:
                continue
            diff = np.array(data[:length], dtype=float) - base
            max_dev = float(np.max(np.abs(diff)))
            std_dev = float(np.std(diff))
            status = "稳定"
            if drift_limit > 0 and max_dev > drift_limit:
                status = "疑似漂移"
                unstable.append(name)
            messages.append(f"{name}: 最大偏差 {max_dev:.3f} mm, Std {std_dev:.3f} mm, {status}")
        if messages:
            self.last_ref_stability_text = " | ".join(messages)
        else:
            self.last_ref_stability_text = None
        return unstable

    def calc_residuals(self, raw, pred):
        raw_arr = np.array(raw, dtype=float)
        pred_arr = np.array(pred[:len(raw)], dtype=float)
        residuals = raw_arr - pred_arr
        with np.errstate(divide='ignore', invalid='ignore'):
            rel_errors = np.where(raw_arr != 0, np.abs(residuals) / np.abs(raw_arr) * 100, np.nan)
        return residuals, rel_errors

    def residual_check_generic(self, relative_errors, threshold=10.0):
        valid = np.array(relative_errors, dtype=float)
        valid = valid[~np.isnan(valid)]
        if len(valid) == 0:
            return False, "数据不足"
        avg_err = np.mean(valid)
        if avg_err < 1.0: grade = "一级 (优)"
        elif avg_err < 5.0: grade = "二级 (合格)"
        elif avg_err < 10.0: grade = "三级 (勉强)"
        else: grade = "四级 (不合格)"
        msg = f"精度评级: {grade}\n平均相对误差: {avg_err:.2f}%"
        return avg_err <= threshold, msg

    def posterior_variance_check_generic(self, raw, residuals):
        n = len(residuals)
        if n < 1:
            return False, "数据不足", 0.0, 0.0
        raw_arr = np.array(raw, dtype=float)
        s1 = np.std(raw_arr)
        s2 = np.std(residuals)
        C = s2 / s1 if s1 != 0 else 0.0
        res_mean = np.mean(residuals)
        threshold = 0.6745 * s1
        p_count = np.sum(np.abs(residuals - res_mean) < threshold)
        P = p_count / n
        if P > 0.95 and C < 0.35: grade = "一级 (好)"; qualified = True
        elif P > 0.80 and C < 0.50: grade = "二级 (合格)"; qualified = True
        elif P > 0.70 and C < 0.65: grade = "三级 (勉强)"; qualified = True
        else: grade = "四级 (不合格)"; qualified = False
        msg = f"精度评级: {grade}\n后验差比值 C = {C:.4f}\n小误差概率 P = {P:.4f}"
        return qualified, msg, C, P

    def fit_log_model(self, t, y):
        if np.any(t <= 0):
            return None
        x = np.log(t)
        coeffs = np.polyfit(x, y, 1)
        return {"a": coeffs[1], "b": coeffs[0]}

    def predict_log_model(self, t, params):
        return params["a"] + params["b"] * np.log(t)

    def fit_power_model(self, t, y):
        if np.any(t <= 0) or np.any(y <= 0):
            return None
        x = np.log(t)
        yy = np.log(y)
        coeffs = np.polyfit(x, yy, 1)
        return {"a": float(np.exp(coeffs[1])), "b": float(coeffs[0])}

    def predict_power_model(self, t, params):
        return params["a"] * np.power(t, params["b"])

    def fit_exp_model(self, t, y):
        if np.any(y <= 0):
            return None
        coeffs = np.polyfit(t, np.log(y), 1)
        return {"a": float(np.exp(coeffs[1])), "b": float(coeffs[0])}

    def predict_exp_model(self, t, params):
        return params["a"] * np.exp(params["b"] * t)

    def fit_hyperbolic_model(self, t, y):
        if np.any(t <= 0) or np.any(y <= 0):
            return None
        x = 1.0 / t
        yy = 1.0 / y
        coeffs = np.polyfit(x, yy, 1)
        return {"a": coeffs[0], "b": coeffs[1]}

    def predict_hyperbolic_model(self, t, params):
        return t / (params["a"] + params["b"] * t)

    def evaluate_models(self, raw_data, future_periods):
        n = len(raw_data)
        t = np.arange(1, n + 1, dtype=float)
        t_all = np.arange(1, n + future_periods + 1, dtype=float)
        results = []

        gm = GreyModelGM11()
        success, _ = gm.fit(raw_data)
        if success:
            pred = gm.predict(len(t_all))
            residuals, rel_errors = self.calc_residuals(raw_data, pred)
            rmse = float(np.sqrt(np.mean(np.square(residuals))))
            results.append({
                "name": "GM(1,1)",
                "pred": pred,
                "residuals": residuals,
                "rel_errors": rel_errors,
                "rmse": rmse,
                "params": {"a": gm.a, "b": gm.b}
            })

        for name, fit_func, pred_func in [
            ("对数模型", self.fit_log_model, self.predict_log_model),
            ("幂函数模型", self.fit_power_model, self.predict_power_model),
            ("指数模型", self.fit_exp_model, self.predict_exp_model),
            ("双曲线模型", self.fit_hyperbolic_model, self.predict_hyperbolic_model)
        ]:
            try:
                params = fit_func(t, np.array(raw_data, dtype=float))
                if params is None:
                    continue
                pred = pred_func(t_all, params)
                residuals, rel_errors = self.calc_residuals(raw_data, pred)
                rmse = float(np.sqrt(np.mean(np.square(residuals))))
                results.append({
                    "name": name,
                    "pred": pred,
                    "residuals": residuals,
                    "rel_errors": rel_errors,
                    "rmse": rmse,
                    "params": params
                })
            except Exception:
                continue
        return results

    def compute_multi_point_metrics(self, raw_data):
        """计算多点变形指标（基于参考点）"""
        ref_name = self.combo_ref_point.currentText()
        if not ref_name or ref_name not in self.points_data or ref_name == self.current_point:
            return None
        ref_data = self.points_data.get(ref_name, [])
        if not ref_data:
            return None
        n = min(len(raw_data), len(ref_data))
        if n < 2:
            return None
        cur_last = raw_data[n - 1]
        ref_last = ref_data[n - 1]
        diff = cur_last - ref_last
        distance = abs(self.point_positions.get(self.current_point, 0.0) - self.point_positions.get(ref_name, 0.0))
        gradient = diff / distance if distance > 0 else None  # mm/m
        tilt_deg = np.degrees(np.arctan2(diff, distance)) if distance > 0 else None
        angular_distortion = (distance * 1000.0 / abs(diff)) if diff != 0 else None  # 1/ratio
        diff_series = np.array(raw_data[:n]) - np.array(ref_data[:n])
        max_abs_diff = float(np.max(np.abs(diff_series)))
        return {
            "diff": diff,
            "gradient": gradient,
            "tilt_deg": tilt_deg,
            "angular_distortion": angular_distortion,
            "max_abs_diff": max_abs_diff,
            "ref_name": ref_name,
            "distance": distance
        }

    def compute_metrics_for_point(self, point_name, raw_data):
        ref_name = self.combo_ref_point.currentText()
        if not ref_name or ref_name not in self.points_data or ref_name == point_name:
            return None
        ref_data = self.points_data.get(ref_name, [])
        if not ref_data:
            return None
        n = min(len(raw_data), len(ref_data))
        if n < 2:
            return None
        cur_last = raw_data[n - 1]
        ref_last = ref_data[n - 1]
        diff = cur_last - ref_last
        distance = abs(self.point_positions.get(point_name, 0.0) - self.point_positions.get(ref_name, 0.0))
        gradient = diff / distance if distance > 0 else None
        tilt_deg = np.degrees(np.arctan2(diff, distance)) if distance > 0 else None
        angular_distortion = (distance * 1000.0 / abs(diff)) if diff != 0 else None
        diff_series = np.array(raw_data[:n]) - np.array(ref_data[:n])
        max_abs_diff = float(np.max(np.abs(diff_series)))
        return {
            "diff": diff,
            "gradient": gradient,
            "tilt_deg": tilt_deg,
            "angular_distortion": angular_distortion,
            "max_abs_diff": max_abs_diff,
            "ref_name": ref_name,
            "distance": distance
        }

    def assess_risk_stability(self, raw_data, metrics):
        try:
            stable_window = int(self.input_stable_window.text())
        except ValueError:
            stable_window = 3
        try:
            stable_threshold = float(self.input_stable_threshold.text())
        except ValueError:
            stable_threshold = 0.5

        stability_status = "--"
        if len(raw_data) >= stable_window + 1 and stable_window > 0:
            rates = np.diff(raw_data)
            stable = np.all(np.abs(rates[-stable_window:]) <= stable_threshold)
            stability_status = "趋稳" if stable else "未趋稳"

        try:
            total_limit = float(self.input_total_limit.text())
        except ValueError:
            total_limit = 30.0
        try:
            diff_limit = float(self.input_diff_limit.text())
        except ValueError:
            diff_limit = 15.0
        try:
            tilt_limit = float(self.input_tilt_limit.text())
        except ValueError:
            tilt_limit = 2.0

        cur_last = raw_data[-1]
        risk = "绿色"
        red_flag = False
        if total_limit > 0 and abs(cur_last) > total_limit:
            red_flag = True
        if metrics and diff_limit > 0 and abs(metrics["diff"]) > diff_limit:
            red_flag = True
        if metrics and tilt_limit > 0 and metrics["gradient"] is not None and abs(metrics["gradient"]) > tilt_limit:
            red_flag = True

        if red_flag:
            risk = "红色"
        else:
            yellow_flag = False
            if total_limit > 0 and abs(cur_last) > 0.8 * total_limit:
                yellow_flag = True
            if metrics and diff_limit > 0 and abs(metrics["diff"]) > 0.8 * diff_limit:
                yellow_flag = True
            if metrics and tilt_limit > 0 and metrics["gradient"] is not None and abs(metrics["gradient"]) > 0.8 * tilt_limit:
                yellow_flag = True
            if stability_status == "未趋稳":
                yellow_flag = True
            if yellow_flag:
                risk = "黄色"
        return risk, stability_status

    def grade_deformation(self, metrics):
        if not metrics:
            return None
        try:
            diff_limit = float(self.input_diff_limit.text())
        except ValueError:
            diff_limit = 15.0
        try:
            tilt_limit = float(self.input_tilt_limit.text())
        except ValueError:
            tilt_limit = 2.0
        try:
            ang_limit = float(self.input_ang_limit.text())
        except ValueError:
            ang_limit = 500.0

        def grade_linear(value, limit):
            if limit <= 0:
                return "未评定"
            if abs(value) >= limit:
                return "红色"
            if abs(value) >= 0.8 * limit:
                return "黄色"
            return "绿色"

        grade_diff = grade_linear(metrics.get("diff", 0.0), diff_limit)
        grade_tilt = grade_linear(metrics.get("gradient", 0.0) if metrics.get("gradient") is not None else 0.0, tilt_limit)
        # 角变形为 1/ratio，数值越小越危险
        ang_val = metrics.get("angular_distortion")
        if ang_val is None or ang_limit <= 0:
            grade_ang = "未评定"
        else:
            if ang_val <= ang_limit:
                grade_ang = "红色"
            elif ang_val <= ang_limit * 1.25:
                grade_ang = "黄色"
            else:
                grade_ang = "绿色"

        return {"diff": grade_diff, "tilt": grade_tilt, "ang": grade_ang}

    def update_accuracy_evaluation(self):
        try:
            inst = float(self.input_inst_precision.text())
        except ValueError:
            inst = 1.0
        try:
            rounds = int(self.input_rounds.text())
        except ValueError:
            rounds = 2
        try:
            obs_std = float(self.input_obs_std.text())
        except ValueError:
            obs_std = 1.0

        if rounds <= 0:
            rounds = 1
        sigma_mean = inst / np.sqrt(rounds)
        allowed = sigma_mean * 2.0
        qualified = obs_std <= allowed
        status = "合格" if qualified else "不合格"
        self.lbl_acc_eval.setText(f"评定：{status} (σ={obs_std:.3f}, 允许≤{allowed:.3f})")
        return {"inst": inst, "rounds": rounds, "obs_std": obs_std, "allowed": allowed, "status": status}

    def compute_resurvey_plan(self):
        rows = []
        names = list(self.points_data.keys())
        positions = {n: self.point_positions.get(n, 0.0) for n in names}
        for name in names:
            res = self.point_results.get(name, {})
            risk = res.get("risk", "--")
            stability = res.get("stability", "--")
            pos = positions.get(name, 0.0)
            distances = [abs(pos - positions[o]) for o in names if o != name]
            nearest = min(distances) if distances else None

            # 基础频率
            if risk == "红色" or stability == "未趋稳":
                days = 3
            elif risk == "黄色":
                days = 7
            else:
                days = 14

            # 间距调整
            if nearest is not None:
                if nearest > 20:
                    days = max(1, int(days * 0.7))
                elif nearest < 5:
                    days = int(days * 1.2)

            freq = f"每 {days} 天"
            rows.append([name, f"{nearest:.2f}" if nearest is not None else "--", freq])
        return rows

    def update_resurvey_table(self):
        rows = self.compute_resurvey_plan()
        self.table_resurvey.setRowCount(0)
        for i, row in enumerate(rows):
            self.table_resurvey.insertRow(i)
            for j, val in enumerate(row):
                self.table_resurvey.setItem(i, j, QTableWidgetItem(str(val)))

    def predict_stable_date(self, pred_vals, date_list, threshold, window):
        if pred_vals is None or len(pred_vals) == 0 or window <= 0:
            return None
        rates = np.diff(pred_vals)
        if len(rates) < window:
            return None
        for i in range(window - 1, len(rates)):
            window_rates = rates[i - window + 1:i + 1]
            if np.all(np.abs(window_rates) <= threshold):
                idx = min(i + 1, len(date_list) - 1)
                return date_list[idx]
        return None

    def compute_multi_point_analysis(self):
        if not self.point_results:
            return None
        valid = [(n, r) for n, r in self.point_results.items() if not r.get("error") and r.get("raw_used")]
        if not valid:
            return None
        min_len = min(len(r["raw_used"]) for _, r in valid)
        if min_len <= 0:
            return None
        curves = np.array([r["raw_used"][:min_len] for _, r in valid], dtype=float)
        mean_curve = np.mean(curves, axis=0).tolist()
        last_mean = mean_curve[-1]
        max_diff_point = None
        max_diff_val = None
        for name, r in valid:
            val = r["raw_used"][min_len - 1]
            diff = abs(val - last_mean)
            if max_diff_val is None or diff > max_diff_val:
                max_diff_val = diff
                max_diff_point = name

        outlier_map = {}
        for name, r in valid:
            for idx in r.get("outliers", []):
                outlier_map.setdefault(idx, []).append(name)
        propagation = [idx + 1 for idx, names in outlier_map.items() if len(names) >= 2]

        return {
            "mean_curve": mean_curve,
            "max_diff_point": max_diff_point,
            "max_diff_value": max_diff_val,
            "propagation_periods": sorted(propagation),
            "propagation_count": len(propagation)
        }

    def update_multi_point_analysis_ui(self):
        if not self.multi_point_analysis:
            self.lbl_mean_curve.setText("群点平均曲线：--")
            self.lbl_max_diff.setText("最大差异点：--")
            self.lbl_outlier_spread.setText("异常传播：--")
            return
        m = self.multi_point_analysis
        self.lbl_mean_curve.setText(f"群点平均曲线：已计算 (期数 {len(m['mean_curve'])})")
        if m["max_diff_point"] is not None:
            self.lbl_max_diff.setText(f"最大差异点：{m['max_diff_point']} (差值 {m['max_diff_value']:.3f} mm)")
        else:
            self.lbl_max_diff.setText("最大差异点：--")
        if m["propagation_count"] > 0:
            periods = ", ".join([str(p) for p in m["propagation_periods"]])
            self.lbl_outlier_spread.setText(f"异常传播：{m['propagation_count']} 期 ({periods})")
        else:
            self.lbl_outlier_spread.setText("异常传播：未发现")

    def update_multi_point_metrics(self):
        """刷新多点变形指标面板"""
        raw_data = self.get_table_data(allow_empty=True)
        if raw_data is None:
            return
        if not raw_data:
            self.lbl_diff_settlement.setText("差异沉降：-- mm")
            self.lbl_tilt.setText("倾斜率：-- mm/m")
            self.lbl_ang_distortion.setText("角变形：--")
            self.lbl_stability.setText("稳定性：--")
            self.lbl_risk.setText("风险等级：--")
            return

        metrics = self.compute_multi_point_metrics(raw_data)
        if metrics:
            diff = metrics["diff"]
            self.lbl_diff_settlement.setText(f"差异沉降：{diff:+.2f} mm (参考点 {metrics['ref_name']})")
            if metrics["gradient"] is not None:
                self.lbl_tilt.setText(f"倾斜率：{metrics['gradient']:.3f} mm/m (倾角 {metrics['tilt_deg']:.3f}°)")
            else:
                self.lbl_tilt.setText("倾斜率：-- mm/m")
            if metrics["angular_distortion"] is not None:
                self.lbl_ang_distortion.setText(f"角变形：1/{metrics['angular_distortion']:.0f}")
            else:
                self.lbl_ang_distortion.setText("角变形：--")
        else:
            self.lbl_diff_settlement.setText("差异沉降：-- mm")
            self.lbl_tilt.setText("倾斜率：-- mm/m")
            self.lbl_ang_distortion.setText("角变形：--")

        # 稳定性判定
        try:
            stable_window = int(self.input_stable_window.text())
        except ValueError:
            stable_window = 3
        try:
            stable_threshold = float(self.input_stable_threshold.text())
        except ValueError:
            stable_threshold = 0.5

        stability_status = "--"
        if len(raw_data) >= stable_window + 1 and stable_window > 0:
            rates = np.diff(raw_data)
            stable = np.all(np.abs(rates[-stable_window:]) <= stable_threshold)
            stability_status = "趋稳" if stable else "未趋稳"
        self.lbl_stability.setText(f"稳定性：{stability_status}")
        self.last_stability_text = stability_status

        # 风险等级评估
        try:
            total_limit = float(self.input_total_limit.text())
        except ValueError:
            total_limit = 30.0
        try:
            diff_limit = float(self.input_diff_limit.text())
        except ValueError:
            diff_limit = 15.0
        try:
            tilt_limit = float(self.input_tilt_limit.text())
        except ValueError:
            tilt_limit = 2.0

        cur_last = raw_data[-1]
        risk = "绿色"
        red_flag = False
        if total_limit > 0 and abs(cur_last) > total_limit:
            red_flag = True
        if metrics and diff_limit > 0 and abs(metrics["diff"]) > diff_limit:
            red_flag = True
        if metrics and tilt_limit > 0 and metrics["gradient"] is not None and abs(metrics["gradient"]) > tilt_limit:
            red_flag = True

        if red_flag:
            risk = "红色"
        else:
            yellow_flag = False
            if total_limit > 0 and abs(cur_last) > 0.8 * total_limit:
                yellow_flag = True
            if metrics and diff_limit > 0 and abs(metrics["diff"]) > 0.8 * diff_limit:
                yellow_flag = True
            if metrics and tilt_limit > 0 and metrics["gradient"] is not None and abs(metrics["gradient"]) > 0.8 * tilt_limit:
                yellow_flag = True
            if stability_status == "未趋稳":
                yellow_flag = True
            if yellow_flag:
                risk = "黄色"
        self.lbl_risk.setText(f"风险等级：{risk}")
        self.last_risk_text = risk

        grades = self.grade_deformation(metrics)
        if grades:
            self.lbl_grade_diff.setText(f"差异沉降等级：{grades['diff']}")
            self.lbl_grade_tilt.setText(f"倾斜等级：{grades['tilt']}")
            self.lbl_grade_ang.setText(f"角变形等级：{grades['ang']}")
            self.last_deform_grades = grades
        else:
            self.lbl_grade_diff.setText("差异沉降等级：--")
            self.lbl_grade_tilt.setText("倾斜等级：--")
            self.lbl_grade_ang.setText("角变形等级：--")
            self.last_deform_grades = None
    def run_prediction(self):
        self.save_current_point_data()
        self.compute_all_points()
        if not self.point_results:
            QMessageBox.warning(self, "数据不足", "当前没有可计算的测点数据。")
            return

        if self.chk_apply_ref.isChecked() and self.last_unstable_refs:
            QMessageBox.warning(self, "基准点稳定性提示", "以下基准点疑似漂移：\n" + "\n".join(self.last_unstable_refs))

        current = self.current_point
        if current in self.point_results and not self.point_results[current].get("error"):
            self.display_point_results(current)
        else:
            for name, res in self.point_results.items():
                if not res.get("error"):
                    self.display_point_results(name)
                    break
        self.update_accuracy_evaluation()
        self.update_resurvey_table()
        self.multi_point_analysis = self.compute_multi_point_analysis()
        self.update_multi_point_analysis_ui()

    def export_to_excel(self):
        """将预测结果和精度评定导出到 Excel 文件"""
        if not HAS_OPENPYXL:
            QMessageBox.critical(self, "缺少依赖", "未检测到 openpyxl 库，无法导出 Excel。\n请先安装: pip install openpyxl")
            return

        if not self.point_results:
            self.compute_all_points()
        if not self.point_results:
            QMessageBox.warning(self, "提示", "请先执行计算与预测。")
            return

        fp, _ = QFileDialog.getSaveFileName(self, "导出 Excel", "settlement_prediction.xlsx", "Excel Files (*.xlsx)")
        if not fp:
            return

        try:
            # 1. 创建工作簿
            wb = Workbook()
            ws_summary = wb.active
            ws_summary.title = "汇总"
            ws_summary.append(["测点", "模型", "RMSE", "MAE", "风险等级", "稳定性", "异常点数"])
            ws_full = wb.create_sheet(title="完整计算序列")
            ws_full.append([
                "测点", "期数", "原始值", "处理值", "预测值", "残差", "相对误差(%)",
                "速率", "施工阶段", "异常标记",
                "多项式拟合", "滚动预测", "卡尔曼滤波", "移动平均", "指数平滑",
                "置信区间下界", "置信区间上界",
                "差异沉降", "倾斜率", "基准修正", "异常处理"
            ])
            ws_before_after = wb.create_sheet(title="处理前后序列")
            ws_before_after.append(["测点", "期数", "处理前", "处理后", "异常处理", "基准修正"])

            current = self.current_point
            image_files = []
            try:
                from openpyxl.drawing.image import Image as XLImage
            except Exception:
                XLImage = None

            for name, res in self.point_results.items():
                if res.get("error"):
                    continue
                ws = wb.create_sheet(title=name)
                raw_used = res["raw_used"]
                raw_original = res["raw_original"]
                pred = res["pred"]
                resid = res["residuals"]
                rel_err = res["rel_errors"]
                stages = self.points_stage.get(name, [])
                outliers = set(res.get("outliers", []))
                poly_vals, rolling_vals, kalman_vals, ma_vals, ema_vals = self.compute_optional_layers(raw_used, total_periods=len(pred))
                conf_bounds = res.get("conf_bounds")
                corrected = res.get("corrected", False)
                outlier_strategy = res.get("outlier_strategy", "不处理")

                ws.append(["期数", "处理值", "预测值", "残差", "相对误差(%)", "施工阶段"])
                for i in range(len(pred)):
                    real_val = raw_used[i] if i < len(raw_used) else ""
                    stage_val = stages[i] if i < len(stages) else ""
                    row = [
                        i + 1,
                        real_val,
                        pred[i],
                        resid[i] if i < len(raw_used) else "",
                        rel_err[i] if i < len(raw_used) else "",
                        stage_val
                    ]
                    ws.append(row)

                metrics_series = self.compute_metrics_for_point(name, raw_original)
                diff_series = None
                tilt_series = None
                if metrics_series and metrics_series.get("ref_name"):
                    ref_data = self.points_data.get(metrics_series["ref_name"], [])
                    n = min(len(raw_original), len(ref_data))
                    if n > 0:
                        diff_series = (np.array(raw_original[:n]) - np.array(ref_data[:n])).tolist()
                        distance = abs(self.point_positions.get(name, 0.0) - self.point_positions.get(metrics_series["ref_name"], 0.0))
                        if distance > 0:
                            tilt_series = (np.array(diff_series) / distance).tolist()

                for i in range(len(pred)):
                    orig_val = raw_original[i] if i < len(raw_original) else ""
                    real_val = raw_used[i] if i < len(raw_used) else ""
                    prev = raw_used[i - 1] if i - 1 >= 0 and i - 1 < len(raw_used) else ""
                    rate = real_val - prev if real_val != "" and prev != "" else ""
                    stage_val = stages[i] if i < len(stages) else ""
                    out_flag = "异常" if i in outliers else ""
                    poly_val = poly_vals[i] if poly_vals is not None and i < len(poly_vals) else ""
                    rolling_val = rolling_vals[i] if rolling_vals is not None and i < len(rolling_vals) else ""
                    kalman_val = kalman_vals[i] if kalman_vals is not None and i < len(kalman_vals) else ""
                    ma_val = ma_vals[i] if ma_vals is not None and i < len(ma_vals) else ""
                    ema_val = ema_vals[i] if ema_vals is not None and i < len(ema_vals) else ""
                    conf_low = conf_bounds[0][i] if conf_bounds is not None and i < len(conf_bounds[0]) else ""
                    conf_high = conf_bounds[1][i] if conf_bounds is not None and i < len(conf_bounds[1]) else ""
                    diff_val = diff_series[i] if diff_series is not None and i < len(diff_series) else ""
                    tilt_val = tilt_series[i] if tilt_series is not None and i < len(tilt_series) else ""
                    ws_full.append([
                        name,
                        i + 1,
                        orig_val,
                        real_val,
                        pred[i],
                        resid[i] if i < len(raw_used) else "",
                        rel_err[i] if i < len(raw_used) else "",
                        rate,
                        stage_val,
                        out_flag,
                        poly_val,
                        rolling_val,
                        kalman_val,
                        ma_val,
                        ema_val,
                        conf_low,
                        conf_high,
                        diff_val,
                        tilt_val,
                        "是" if corrected else "否",
                        outlier_strategy
                    ])

                for i in range(len(raw_original)):
                    before_val = raw_original[i] if i < len(raw_original) else ""
                    after_val = raw_used[i] if i < len(raw_used) else ""
                    ws_before_after.append([
                        name,
                        i + 1,
                        before_val,
                        after_val,
                        outlier_strategy,
                        "是" if corrected else "否"
                    ])

                ws.append([])
                ws.append(["模型", res.get("model_name", "")])
                ws.append(["RMSE", res.get("rmse", "")])
                ws.append(["MAE", res.get("mae", "")])
                ws.append(["风险等级", res.get("risk", "")])
                ws.append(["稳定性", res.get("stability", "")])
                if res.get("stable_date"):
                    ws.append(["预计趋稳日期", res["stable_date"].strftime("%Y-%m-%d")])
                if res.get("change_index") is not None:
                    ws.append(["工况变化点", f"第 {res['change_index'] + 1} 期"])
                if res.get("grades"):
                    g = res["grades"]
                    ws.append(["差异沉降等级", g.get("diff", "")])
                    ws.append(["倾斜等级", g.get("tilt", "")])
                    ws.append(["角变形等级", g.get("ang", "")])

                if res.get("stage_stats"):
                    ws.append([])
                    ws.append(["施工阶段对比"])
                    ws.append(["阶段", "样本数", "平均速率", "最大速率"])
                    for s in res["stage_stats"]:
                        ws.append([s["stage"], s["count"], s["avg_rate"], s["max_rate"]])
                if res.get("stage_deep"):
                    ws.append([])
                    ws.append(["施工阶段深化"])
                    ws.append(["阶段", "变化率", "影响指数", "敏感度排名"])
                    for s in res["stage_deep"]:
                        ws.append([s["stage"], s["change_rate"] if s["change_rate"] is not None else "", s["impact"], s["rank"]])

                if outliers:
                    ws.append([])
                    ws.append(["异常点期次", ", ".join([str(i + 1) for i in sorted(outliers)])])

                ws_summary.append([
                    name,
                    res.get("model_name", ""),
                    res.get("rmse", ""),
                    res.get("mae", ""),
                    res.get("risk", ""),
                    res.get("stability", ""),
                    len(res.get("outliers", []))
                ])

                if XLImage is not None:
                    try:
                        self.display_point_results(name)
                        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tmp.close()
                        self.canvas.figure.savefig(tmp.name, dpi=200, bbox_inches='tight')
                        image_files.append(tmp.name)
                        img = XLImage(tmp.name)
                        img.anchor = "H2"
                        ws.add_image(img)
                    except Exception:
                        pass
            if current and current in self.point_results and not self.point_results[current].get("error"):
                self.display_point_results(current)

            # 观测精度评定与复测频率建议
            ws_plan = wb.create_sheet(title="复测频率建议")
            ws_plan.append(["测点", "最近间距(m)", "建议频率"])
            for row in self.compute_resurvey_plan():
                ws_plan.append(row)

            ws_acc = wb.create_sheet(title="观测精度评定")
            acc = self.update_accuracy_evaluation()
            ws_acc.append(["仪器精度(mm)", acc["inst"]])
            ws_acc.append(["测回数", acc["rounds"]])
            ws_acc.append(["观测标准差(mm)", acc["obs_std"]])
            ws_acc.append(["允许标准差(mm)", acc["allowed"]])
            ws_acc.append(["评定", acc["status"]])

            ws_coop = wb.create_sheet(title="协同分析")
            if self.multi_point_analysis is None:
                self.multi_point_analysis = self.compute_multi_point_analysis()
            if self.multi_point_analysis:
                m = self.multi_point_analysis
                ws_coop.append(["群点平均曲线长度", len(m["mean_curve"])])
                ws_coop.append(["期数", "群点平均值"])
                for i, val in enumerate(m["mean_curve"], start=1):
                    ws_coop.append([i, val])
                if m["max_diff_point"] is not None:
                    ws_coop.append(["最大差异点", m["max_diff_point"]])
                    ws_coop.append(["最大差异值(mm)", m["max_diff_value"]])
                if m["propagation_count"] > 0:
                    ws_coop.append(["异常传播期次", ", ".join([str(p) for p in m["propagation_periods"]])])
                else:
                    ws_coop.append(["异常传播期次", "未发现"])

            # 6. 保存文件
            wb.save(fp)
            QMessageBox.information(self, "导出成功", f"数据已导出到:\n{fp}")

        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出到 Excel 失败:\n{e}")
        finally:
            try:
                for img_path in image_files:
                    if os.path.exists(img_path):
                        os.remove(img_path)
            except Exception:
                pass

    def export_word_report(self):
        """生成 Word 格式的分析报告"""
        if not HAS_DOCX:
            QMessageBox.critical(self, "缺少依赖", "未检测到 python-docx 库，无法生成 Word 报告。\n请先安装: pip install python-docx")
            return

        if not self.point_results:
            self.compute_all_points()
        if not self.point_results:
            QMessageBox.warning(self, "提示", "请先执行计算与预测。")
            return

        fp, _ = QFileDialog.getSaveFileName(self, "保存报告", f"沉降监测报告_{datetime.now().strftime('%Y%m%d')}.docx", "Word Files (*.docx)")
        if not fp:
            return

        try:
            doc = Document()
            
            # 1. 标题
            heading = doc.add_heading('沉降监测与变形预测报告', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"预测模型: 多测点自动处理")

            doc.add_heading('1. 监测简报', level=1)
            brief_items, brief_suggestions = self.build_monitor_brief_all()
            for line in brief_items:
                doc.add_paragraph(line)
            for line in brief_suggestions:
                doc.add_paragraph(line)

            current = self.current_point
            for name, res in self.point_results.items():
                if res.get("error"):
                    continue
                doc.add_heading(f'测点 {name}', level=1)

                p = doc.add_paragraph()
                p.add_run(f"模型: {res.get('model_name','')}\n")
                p.add_run(f"RMSE: {res.get('rmse',0):.4f}, MAE: {res.get('mae',0):.4f}\n")
                p.add_run(f"风险等级: {res.get('risk','--')} | 稳定性: {res.get('stability','--')}\n")
                if res.get("stable_date"):
                    p.add_run(f"预计趋稳日期: {res['stable_date'].strftime('%Y-%m-%d')}\n")

                if res.get("change_index") is not None:
                    doc.add_paragraph(f"工况变化点: 第 {res['change_index'] + 1} 期")
                if res.get("grades"):
                    g = res["grades"]
                    doc.add_paragraph(f"差异沉降等级: {g.get('diff','--')} | 倾斜等级: {g.get('tilt','--')} | 角变形等级: {g.get('ang','--')}")

                if res.get("stage_stats"):
                    stage_table = doc.add_table(rows=1, cols=4)
                    stage_table.style = 'Table Grid'
                    hdr = stage_table.rows[0].cells
                    hdr[0].text = '阶段'
                    hdr[1].text = '样本数'
                    hdr[2].text = '平均速率'
                    hdr[3].text = '最大速率'
                    for s in res["stage_stats"]:
                        row = stage_table.add_row().cells
                        row[0].text = s['stage']
                        row[1].text = str(s['count'])
                        row[2].text = f"{s['avg_rate']:.3f}"
                        row[3].text = f"{s['max_rate']:.3f}"
                if res.get("stage_deep"):
                    deep_table = doc.add_table(rows=1, cols=4)
                    deep_table.style = 'Table Grid'
                    hdr = deep_table.rows[0].cells
                    hdr[0].text = '阶段'
                    hdr[1].text = '变化率'
                    hdr[2].text = '影响指数'
                    hdr[3].text = '敏感度排名'
                    for s in res["stage_deep"]:
                        row = deep_table.add_row().cells
                        row[0].text = s['stage']
                        row[1].text = f"{s['change_rate']:.3f}" if s["change_rate"] is not None else "--"
                        row[2].text = f"{s['impact']:.3f}"
                        row[3].text = str(s['rank'])

                try:
                    self.display_point_results(name)
                    img_stream = io.BytesIO()
                    self.canvas.figure.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
                    img_stream.seek(0)
                    doc.add_picture(img_stream, width=Inches(6.0))
                except Exception:
                    pass

                raw = res["raw_used"]
                pred = res["pred"]
                resid = res["residuals"]
                rel_err = res["rel_errors"]
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '期数'
                hdr_cells[1].text = '处理值 (mm)'
                hdr_cells[2].text = '预测值 (mm)'
                hdr_cells[3].text = '残差 (mm)'
                hdr_cells[4].text = '相对误差 (%)'
                hdr_cells[5].text = '施工阶段'

                stages = self.points_stage.get(name, [])
                for i in range(len(pred)):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i + 1)
                    if i < len(raw):
                        row_cells[1].text = f"{raw[i]:.2f}"
                        row_cells[3].text = f"{resid[i]:.2f}"
                        row_cells[4].text = f"{rel_err[i]:.2f}"
                    else:
                        row_cells[1].text = "-"
                        row_cells[3].text = "-"
                        row_cells[4].text = "-"
                    row_cells[2].text = f"{pred[i]:.2f}"
                    row_cells[5].text = stages[i] if i < len(stages) else ""

            doc.add_heading('复测频率建议表', level=1)
            resurvey_table = doc.add_table(rows=1, cols=3)
            resurvey_table.style = 'Table Grid'
            hdr = resurvey_table.rows[0].cells
            hdr[0].text = '测点'
            hdr[1].text = '最近间距(m)'
            hdr[2].text = '建议频率'
            for row in self.compute_resurvey_plan():
                r = resurvey_table.add_row().cells
                r[0].text = str(row[0])
                r[1].text = str(row[1])
                r[2].text = str(row[2])

            doc.add_heading('观测精度评定', level=1)
            acc = self.update_accuracy_evaluation()
            doc.add_paragraph(f"仪器精度: {acc['inst']:.3f} mm | 测回数: {acc['rounds']} | 观测标准差: {acc['obs_std']:.3f} mm")
            doc.add_paragraph(f"允许标准差: {acc['allowed']:.3f} mm | 评定: {acc['status']}")

            doc.add_heading('多测点协同分析', level=1)
            if self.multi_point_analysis:
                m = self.multi_point_analysis
                if m["max_diff_point"] is not None:
                    doc.add_paragraph(f"最大差异点: {m['max_diff_point']} (差值 {m['max_diff_value']:.3f} mm)")
                if m["propagation_count"] > 0:
                    periods = ", ".join([str(p) for p in m["propagation_periods"]])
                    doc.add_paragraph(f"异常传播期次: {periods}")
                else:
                    doc.add_paragraph("异常传播: 未发现")

            doc.add_heading('2. AI 总结', level=1)
            brief_items, brief_suggestions = self.build_monitor_brief_all()
            summary_text, err = self.generate_ai_summary_all(brief_items, brief_suggestions)
            if summary_text:
                doc.add_paragraph(summary_text)
            else:
                doc.add_paragraph(f"AI 总结生成失败: {err if err else '未配置 API Key'}")

            if current and current in self.point_results and not self.point_results[current].get("error"):
                self.display_point_results(current)

            doc.save(fp)
            QMessageBox.information(self, "导出成功", f"报告已生成:\n{fp}")

        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"生成报告时发生错误:\n{e}")

    def export_image(self):
        """导出当前图表为图片文件"""
        if not self.point_results:
             QMessageBox.warning(self, "提示", "请先执行计算生成图表后再导出。")
             return

        folder = QFileDialog.getExistingDirectory(self, "选择导出文件夹")
        if not folder:
            return
        current = self.current_point
        try:
            for name, res in self.point_results.items():
                if res.get("error"):
                    continue
                self.display_point_results(name)
                fp = os.path.join(folder, f"{name}_analysis.png")
                self.canvas.figure.savefig(fp, dpi=300, bbox_inches='tight')
                if res.get("outlier_applied"):
                    fig_cmp = self.build_outlier_compare_figure(res.get("raw_original", []), res.get("raw_used", []), name)
                    if fig_cmp is not None:
                        fp_cmp = os.path.join(folder, f"{name}_outlier_compare.png")
                        fig_cmp.savefig(fp_cmp, dpi=300, bbox_inches='tight')
                        plt.close(fig_cmp)
            QMessageBox.information(self, "导出成功", f"图表已保存至:\n{folder}")
        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"保存图片时发生错误:\n{e}")
        finally:
            if current and current in self.point_results and not self.point_results[current].get("error"):
                self.display_point_results(current)

    def update_log(self, raw, pred, resid, rel_err, grade_msg="", grade_msg_cp="", rate_threshold=2.0, date_list=None):
        self.log_text.clear()
        model_name = self.last_model_name or "GM(1,1)"
        self.log_text.append(f"=== {model_name} 模型计算报告 ===")
        if self.current_point:
            self.log_text.append(f"当前测点: {self.current_point}")
        if model_name == "GM(1,1)" and self.last_model_params:
            self.log_text.append(f"发展系数 a = {self.last_model_params.get('a', 0.0):.6f} (反映发展趋势)")
            self.log_text.append(f"灰作用量 b = {self.last_model_params.get('b', 0.0):.6f}")
        elif self.last_model_params:
            self.log_text.append(f"模型参数: {self.last_model_params}")
        if self.last_corrected:
            self.log_text.append("基准漂移修正: 已应用到当前测点")
        if self.last_model_rmse is not None:
            self.log_text.append(f"RMSE: {self.last_model_rmse:.4f} | MAE: {self.last_model_mae:.4f}")
        
        # 简单的模型评价
        if grade_msg:
            self.log_text.append(f"--- 精度检验 ---\n{grade_msg}")
        
        if grade_msg_cp:
            self.log_text.append(f"--- 后验差检验 (C/P) ---\n{grade_msg_cp}")
            
        if model_name == "GM(1,1)" and self.last_model_params:
            a_val = abs(self.last_model_params.get("a", 0.0))
            if a_val < 0.3:
                self.log_text.append("评价: 模型适用于中长期预测 (a < 0.3)")
            elif a_val < 0.5:
                self.log_text.append("评价: 模型适用于短期预测 (0.3 < a < 0.5)")
            else:
                self.log_text.append("警告: 发展系数较大，预测需谨慎！")
        elif model_name == "指数模型" and self.last_model_params:
            b_val = self.last_model_params.get("b", 0.0)
            if b_val < 0:
                self.log_text.append("参数说明：指数模型趋于 0（稳定值）")
            else:
                self.log_text.append("参数说明：指数模型为发散趋势")
        elif model_name == "双曲线模型" and self.last_model_params:
            b_val = self.last_model_params.get("b", 0.0)
            if b_val != 0:
                stable = 1.0 / b_val
                self.log_text.append(f"参数说明：双曲线稳定值约为 {stable:.4f}")
        if self.last_change_index is not None:
            self.log_text.append(f"工况变化点: 第 {self.last_change_index + 1} 期")
        if self.chk_apply_ref.isChecked():
            self.log_text.append("基准漂移修正: 已启用")
        if self.last_ref_stability_text:
            self.log_text.append(f"基准稳定性: {self.last_ref_stability_text}")

        if self.last_stage_stats:
            self.log_text.append("\n--- 施工阶段对比 ---")
            self.log_text.append(f"{'阶段':<8}\t{'样本数':<6}\t{'平均速率':<10}\t{'最大速率':<10}")
            for s in self.last_stage_stats:
                self.log_text.append(f"{s['stage']:<8}\t{s['count']:<6}\t{s['avg_rate']:<10.3f}\t{s['max_rate']:<10.3f}")
        if self.last_stage_deep:
            self.log_text.append("\n--- 施工阶段深化 ---")
            self.log_text.append(f"{'阶段':<8}\t{'变化率':<10}\t{'影响指数':<10}\t{'敏感度排名':<10}")
            for s in self.last_stage_deep:
                cr = f"{s['change_rate']:.3f}" if s["change_rate"] is not None else "--"
                self.log_text.append(f"{s['stage']:<8}\t{cr:<10}\t{s['impact']:<10.3f}\t{s['rank']:<10}")

        brief_items, brief_suggestions = self.build_monitor_brief()
        self.log_text.append("\n--- 监测简报 ---")
        for line in brief_items:
            self.log_text.append(line)
        for line in brief_suggestions:
            self.log_text.append(line)
        if hasattr(self, "last_stable_date") and self.last_stable_date:
            self.log_text.append(f"预计趋稳日期: {self.last_stable_date.strftime('%Y-%m-%d')}")

        acc = self.update_accuracy_evaluation()
        self.log_text.append("\n--- 观测精度评定 ---")
        self.log_text.append(f"仪器精度: {acc['inst']:.3f} mm | 测回数: {acc['rounds']} | 观测标准差: {acc['obs_std']:.3f} mm")
        self.log_text.append(f"允许标准差: {acc['allowed']:.3f} mm | 评定: {acc['status']}")

        self.log_text.append("\n--- 复测频率建议 ---")
        for row in self.compute_resurvey_plan():
            self.log_text.append(f"{row[0]} | 间距 {row[1]} m | {row[2]}")

        if self.multi_point_analysis:
            m = self.multi_point_analysis
            self.log_text.append("\n--- 多测点协同分析 ---")
            if m["max_diff_point"] is not None:
                self.log_text.append(f"最大差异点: {m['max_diff_point']} (差值 {m['max_diff_value']:.3f} mm)")
            if m["propagation_count"] > 0:
                periods = ", ".join([str(p) for p in m["propagation_periods"]])
                self.log_text.append(f"异常传播期次: {periods}")
            else:
                self.log_text.append("异常传播: 未发现")

        metrics = self.compute_multi_point_metrics(raw)
        if metrics:
            diff = metrics["diff"]
            grad = metrics["gradient"]
            self.log_text.append("\n--- 多点变形指标 ---")
            self.log_text.append(f"参考点: {metrics['ref_name']} (间距 {metrics['distance']:.3f} m)")
            self.log_text.append(f"差异沉降: {diff:+.2f} mm | 最大差异沉降: {metrics['max_abs_diff']:.2f} mm")
            if grad is not None:
                self.log_text.append(f"倾斜率: {grad:.3f} mm/m | 倾角: {metrics['tilt_deg']:.3f}°")
            if metrics["angular_distortion"] is not None:
                self.log_text.append(f"角变形: 1/{metrics['angular_distortion']:.0f}")
            self.log_text.append(self.lbl_stability.text())
            self.log_text.append(self.lbl_risk.text())

        self.log_text.append("\n--- 详细数据 ---")
        self.log_text.append(f"{'日期':<12}\t{'期数':<5}\t{'实测值':<8}\t{'预测值':<8}\t{'残差':<8}\t{'相对误差%'}")
        
        # 历史数据部分
        for i in range(len(raw)):
            d_str = date_list[i].strftime('%Y-%m-%d') if date_list else ""
            self.log_text.append(f"{d_str:<12}\t{i+1:<5}\t{raw[i]:.2f}\t{pred[i]:.2f}\t{resid[i]:.2f}\t{rel_err[i]:.2f}%")
        
        self.log_text.append("\n--- 未来预测 ---")
        # 预测数据部分
        for i in range(len(raw), len(pred)):
            d_str = date_list[i].strftime('%Y-%m-%d') if date_list else ""
            # 简单的速率预警
            rate = pred[i] - pred[i-1]
            warning = " [速率过快!]" if abs(rate) > rate_threshold else ""
            self.log_text.append(f"{d_str:<12}\t{i+1:<5}\t{'--':<8}\t{pred[i]:.2f}\t(增量: {rate:.2f}){warning}")

    def update_plot(self, raw, pred, poly_vals=None, rolling_vals=None, kalman_vals=None, ma_vals=None, rate_threshold=2.0, ema_vals=None, conf_bounds=None, date_list=None, interval_days=1, outlier_indices=None, ref_data=None, ref_name=None, model_name=None, change_index=None, corrected=False, raw_original=None, baseline=None, stage_stats=None, outlier_applied=False, stable_date=None):
        ax1 = self.canvas.ax1
        ax2 = self.canvas.ax2
        ax1.cla() # 清除旧图
        ax2.cla()
        
        # --- 根据主题设置绘图颜色 ---
        bg_color = '#2b2b2b' if self.is_dark_mode else 'white'
        text_color = 'white' if self.is_dark_mode else 'black'
        self.canvas.figure.set_facecolor(bg_color)
        for ax in [ax1, ax2]:
            ax.set_facecolor(bg_color)
            ax.tick_params(colors=text_color)
            for spine in ax.spines.values(): spine.set_color(text_color)
        
        # 使用日期列表作为 X 轴，如果没有则回退到数字
        x_axis = date_list if date_list else np.arange(1, len(pred) + 1)
        
        # 绘制实测数据
        s = self.plot_styles['observed']
        ax1.plot(x_axis[:len(raw)], raw, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                 marker=s['marker'], label='实测沉降值 (Observed)' if not corrected else '实测沉降值 (已修正)', markersize=6)
        if (corrected or outlier_applied) and raw_original is not None:
            label = '实测沉降值 (修正前)' if corrected else '实测沉降值 (处理前)'
            ax1.plot(x_axis[:len(raw_original)], raw_original, color='#9E9E9E', linestyle='--', linewidth=1.2,
                     marker=None, label=label, alpha=0.8)

        # 参考点曲线
        if ref_data and len(ref_data) > 0:
            ref_len = min(len(ref_data), len(x_axis))
            ax1.plot(x_axis[:ref_len], ref_data[:ref_len], color='#888888', linestyle='--', linewidth=1.2,
                     label=f'参考点 {ref_name}', alpha=0.8)
        
        # --- 新增：标记异常值 ---
        if outlier_indices is not None and len(outlier_indices) > 0:
            outlier_x = [x_axis[i] for i in outlier_indices]
            outlier_y = [raw[i] for i in outlier_indices]
            ax1.scatter(outlier_x, outlier_y, 
                        s=150, # 标记大小
                        facecolors='none', # 空心
                        edgecolors='red', # 红色边框
                        linewidths=2,
                        label='异常值 (Outlier)', zorder=10)
        
        # 绘制预测数据 (虚线)
        s = self.plot_styles['predicted']
        ax1.plot(x_axis, pred, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                 marker=s['marker'], label=f'{model_name or "GM(1,1)"} 预测模型 (Predicted)', alpha=0.7)
                 
        # --- 新增：绘制置信区间 (半透明填充) ---
        if conf_bounds is not None:
            s = self.plot_styles['conf_int']
            # 使用 fill_between 填充上下界之间的区域，alpha 设置透明度
            ax1.fill_between(x_axis, conf_bounds[0], conf_bounds[1], color=s['color'], alpha=0.2, label='95% 置信区间')
        
        # 新增：绘制多项式拟合曲线 (绿色虚线)
        if poly_vals is not None:
            s = self.plot_styles['poly']
            poly_order = self.combo_poly_order.currentText()
            if len(poly_vals) != len(x_axis):
                min_len = min(len(poly_vals), len(x_axis))
                ax1.plot(x_axis[:min_len], poly_vals[:min_len], color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                         marker=s['marker'], label=f'{poly_order}阶多项式拟合 (Poly-{poly_order})', alpha=0.6)
            else:
                ax1.plot(x_axis, poly_vals, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                         marker=s['marker'], label=f'{poly_order}阶多项式拟合 (Poly-{poly_order})', alpha=0.6)
            
        # 新增：绘制滚动预测曲线 (洋红色点划线)
        if rolling_vals is not None:
            s = self.plot_styles['rolling']
            ax1.plot(x_axis[:len(raw)], rolling_vals, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                     marker=s['marker'], label='滚动预测检验 (Rolling)', alpha=0.9)
            
        # 新增：绘制卡尔曼滤波曲线 (青色实线)
        if kalman_vals is not None:
            s = self.plot_styles['kalman']
            ax1.plot(x_axis[:len(raw)], kalman_vals, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                     marker=s['marker'], label='卡尔曼滤波 (Kalman)', alpha=0.8)
            
        # 新增：绘制移动平均曲线 (橙色实线)
        if ma_vals is not None:
            s = self.plot_styles['ma']
            # -- 修改：动态更新图例标签 --
            window_size_str = self.input_ma_window.text()
            ax1.plot(x_axis[:len(raw)], ma_vals, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                     marker=s['marker'], label=f'移动平均 (MA-{window_size_str})', alpha=0.8)
            
        # 新增：绘制指数平滑曲线 (紫色实线)
        if ema_vals is not None:
            s = self.plot_styles['ema']
            ax1.plot(x_axis[:len(raw)], ema_vals, color=s['color'], linestyle=s['style'], linewidth=s['width'], 
                     marker=s['marker'], label='指数平滑 (EMA)', alpha=0.8)
        
        # 绘制分界线
        # 如果是日期轴，分界线位置应为最后一个实测点的日期
        vline_pos = x_axis[len(raw)-1] if date_list else len(raw)
        ax1.axvline(x=vline_pos, color='green', linestyle=':', label='当前时间点')
        ax2.axvline(x=vline_pos, color='green', linestyle=':', alpha=0.5)

        if change_index is not None and 0 <= change_index < len(x_axis):
            x_change = x_axis[change_index]
            ax1.axvline(x=x_change, color='orange', linestyle='--', linewidth=1.2, label='工况变化点')
            ax2.axvline(x=x_change, color='orange', linestyle='--', linewidth=1.0, alpha=0.7)

        if stable_date is not None and date_list:
            ax1.axvline(x=stable_date, color='purple', linestyle='--', linewidth=1.2, label='预计趋稳日期')
        
        ax1.set_title("沉降监测 P-T 曲线分析", fontsize=12, color=text_color)
        ax1.set_ylabel("累积沉降量 (mm)", color=text_color)
        ax1.grid(True, linestyle='--', alpha=0.6)
        
        # --- 绘制下方子图：沉降速率 ---
        # 计算速率 (差分)
        rates = np.diff(pred, prepend=pred[0]) # 补一个0保持长度一致
        # 区分实测段和预测段的颜色
        colors = ['#4A90E2'] * len(raw) + ['#E57373'] * (len(pred) - len(raw)) # 蓝色实测，红色预测
        
        # --- 新增：将超出阈值的柱子标为醒目的黄色 ---
        warning_color = '#FFC107' # 醒目的黄色
        for i in range(len(rates)):
            if abs(rates[i]) > rate_threshold:
                colors[i] = warning_color

        # 对于日期轴，bar 的宽度需要调整 (默认是 0.8 天，看起来比较合适)
        bar_width = 0.6 if not date_list else interval_days * 0.6
        ax2.bar(x_axis, rates, color=colors, alpha=0.8, width=bar_width)
        
        ax2.axhline(0, color=text_color, linewidth=0.8, alpha=0.5)
        
        # --- 新增：绘制速率阈值线 ---
        ax2.axhline(rate_threshold, color='red', linestyle='--', linewidth=1.2, alpha=0.9, label=f'速率阈值 (±{rate_threshold})')
        ax2.axhline(-rate_threshold, color='red', linestyle='--', linewidth=1.2, alpha=0.9)

        ax2.set_ylabel("速率 (mm/期)", color=text_color, fontsize=9)
        ax2.set_xlabel("监测日期 (Date)", color=text_color)
        ax2.grid(True, linestyle='--', alpha=0.3)

        if stage_stats:
            try:
                stage_names = [s["stage"] for s in stage_stats]
                stage_vals = [s["avg_rate"] for s in stage_stats]
                ax_inset = inset_axes(ax1, width="35%", height="30%", loc="lower right", borderpad=2)
                ax_inset.bar(range(len(stage_vals)), stage_vals, color="#90CAF9")
                ax_inset.set_title("阶段平均速率", fontsize=8)
                ax_inset.set_xticks(range(len(stage_vals)))
                ax_inset.set_xticklabels(stage_names, rotation=30, fontsize=7)
                ax_inset.tick_params(axis='y', labelsize=7)
            except Exception:
                pass
        
        # --- 日期格式化 ---
        if date_list:
            # 设置日期格式 (例如 2023-01-01)
            date_fmt = mdates.DateFormatter('%Y-%m-%d')
            ax1.xaxis.set_major_formatter(date_fmt)
            ax2.xaxis.set_major_formatter(date_fmt)
            # 自动调整日期标签角度，防止重叠
            self.canvas.figure.autofmt_xdate()
        
        # 隐藏上图的 X 轴标签 (因为共享)
        plt.setp(ax1.get_xticklabels(), visible=False)
        
        leg = ax1.legend(loc='upper left', fontsize=9)
        if self.is_dark_mode:
            for text in leg.get_texts(): text.set_color('white')
            leg.get_frame().set_facecolor('#3c3f41')
            leg.get_frame().set_edgecolor('#555')
            
        # --- 新增：速率图图例 ---
        legend_elements = [
            Patch(facecolor='#4A90E2', alpha=0.8, label='实测'),
            Patch(facecolor='#E57373', alpha=0.8, label='预测'),
            Patch(facecolor='#FFC107', alpha=0.8, label='超限')
        ]
        handles, _ = ax2.get_legend_handles_labels() # 获取阈值线的图例句柄
        leg2 = ax2.legend(handles=handles + legend_elements, loc='upper right', fontsize=8, ncol=4)
        if self.is_dark_mode:
            for text in leg2.get_texts(): text.set_color('white')
            leg2.get_frame().set_facecolor('#3c3f41')
            leg2.get_frame().set_edgecolor('#555')
        
        self.canvas.draw()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    
    # --- 启动画面 (Splash Screen) ---
    # 1. 绘制一个简单的启动图片 (也可以直接加载图片文件: QPixmap("splash.png"))
    splash_pix = QPixmap(600, 350)
    splash_pix.fill(QColor("#2c3e50")) # 深蓝灰色背景
    
    painter = QPainter(splash_pix)
    painter.setPen(QColor("#ecf0f1"))
    painter.setFont(QFont("Microsoft YaHei", 24, QFont.Weight.Bold))
    painter.drawText(splash_pix.rect(), Qt.AlignmentFlag.AlignCenter, "沉降监测与变形预测系统\nLoading...")
    painter.end()

    # 2. 显示启动画面
    splash = QSplashScreen(splash_pix)
    splash.show()
    app.processEvents() # 强制刷新界面，确保画面立即显示
    
    # 3. 模拟加载耗时 (实际项目中这里通常是加载配置文件或数据库)
    time.sleep(1.5)

    window = MainWindow()
    window.show()
    
    # 4. 主窗口显示后关闭启动画面
    splash.finish(window)
    
    sys.exit(app.exec())

